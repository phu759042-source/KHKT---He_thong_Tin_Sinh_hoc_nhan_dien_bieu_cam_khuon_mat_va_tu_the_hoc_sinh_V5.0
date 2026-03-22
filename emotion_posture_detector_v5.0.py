import cv2
import numpy as np
import os
import sys
import time
import socket
from collections import deque, Counter
from threading import Thread, Lock
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
from PIL import Image, ImageDraw, ImageFont, ImageTk
import qrcode
import io
import pyautogui
from pygrabber.dshow_graph import FilterGraph 
import win32gui
import win32clipboard
import win32con
import win32api
from flask import Flask, Response, render_template_string, request
import csv
import datetime
from docx import Document
import requests
import subprocess

SERVER_URL = "https://epd-test.onrender.com/log_incident/"

class_name = False# lớp
SCAN_MODE = "epd_full"# mode gửi server
ZONE_ID = None# đại diện ROI

ROI_STATE_TRACKER = {
    "state": None,
    "start_time": None
}

# Setup đường dẫn
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

icon_path = os.path.join(BASE_DIR, "Emotion + Posture Detector v5.0.ico") 
icon_path_camera = os.path.join(BASE_DIR, "Emotion + Posture Detector v3.0 Camera.ico")
icon_path_fullscreen = os.path.join(BASE_DIR, "Emotion + Posture Detector v3.0 Fullscreen Capture.ico")
font_path = os.path.join(BASE_DIR, "ARIALBD 1.ttf")
# ==================================================

# === THAM SỐ NGƯỠNG THEO BẢNG THAM CHIẾU DOCX ===
THRESHOLD_P_MAX_DEFAULT = 0.01  # Ngưỡng tin cậy chung (τ)
THRESHOLD_DELTA_TOP2_DEFAULT = 0.001 # Ngưỡng nhập nhằng chung (δ)
STABILITY_WINDOW_FRAMES = 1 # Cửa sổ ổn định ngắn hạn W (30 khung)
STABILITY_DOMINANCE_RATIO = 0.8 # Nhãn phải chiếm tối thiểu 80% trong cửa sổ 30 frames để được xác nhận
BAD_POSTURE_WARNING_FRAMES = 600 # Ngưỡng cảnh báo tư thế (khoảng 40 giây ở 15 fps)

# Ngưỡng P_MAX và Delta TOP2 cụ thể theo bảng DOCX
EMOTION_THRESHOLDS = {
    'Giận dữ': {'p_max': 0.01, 'delta': 0.001}, # Giảm ngưỡng P_MAX xuống 1%
    'Ghê sợ': {'p_max': 0.01, 'delta': 0.001},
    'Sợ hãi': {'p_max': 0.01, 'delta': 0.001},
    'Vui vẻ': {'p_max': 0.01, 'delta': 0.001}, 
    'Buồn': {'p_max': 0.01, 'delta': 0.001}, 
    'Bất ngờ': {'p_max': 0.01, 'delta': 0.001},
    'Trung lập': {'p_max': 0.01, 'delta': 0.001}
}
# === HẾT PHẦN KHAI BÁO MỚI ===

# Cảm xúc bất lợi: Buồn, Giận, Sợ hãi
NEGATIVE_EMOTIONS = ['Buồn', 'Giận dữ', 'Sợ hãi'] 

# GLOBALS & KHỞI TẠO CHUNG
latest_frame = None
frame_lock = Lock()
loading_window = None
progress_bar = None
progress_label = None
flask_app = Flask(__name__)
is_running = False 
broadcast_thread = None
current_mode = 'camera' 
detection_thread = None 
thread_lock = Lock()

# CÁC THAY ĐỔI MỚI VỀ LOGGING DỮ LIỆU
SCAN_MIN_DURATION = 1800.0 # Bắt buộc quét tối thiểu 1800 giây ~ 30 phút
DATA_LOGS = []           # Danh sách toàn cục để lưu log dữ liệu
LOG_LOCK = Lock()        # Lock cho việc ghi/đọc DATA_LOGS

history = deque(maxlen=150) # Đưa history về global để truy cập khi dừng
session_start_time = None
session_end_time = None
bad_posture_total_frames = 0
total_detection_frames = 0

# PHẦN VẼ KHUNG
ROI_BOX = None
ROI_DRAWING = False
ROI_ACTIVE = False
roi_start = None
roi_end = None
roi_status_text = "Vẽ khung ROI: TẮT"
roi_status_color = (255, 0, 0)
ROI_IMAGE_PATH = None
ROI_IMAGE_BUFFER = None
roi_emotion_label = None
DISPLAY_SCALE_X = 1.0
DISPLAY_SCALE_Y = 1.0

# FULLSCREEN
INCIDENT_STATE = None
INCIDENT_START_TIME = None
INCIDENT_START_TIME_STR = None

ABNORMAL_THRESHOLD = 6  # giây

ROI_LOGS = []
roi_scan_start_time = None

# GỬI DỮ LIỆU LÊN WEB
ROI_ALERT_HISTORY = deque(maxlen=60)   # lưu 60 giây

log_directory = BASE_DIR # Khai báo biến toàn cục cho thư mục log, mặc định là BASE_DIR

force_exit_no_report = False
# UTILITIES CHUNG

def set_opencv_window_icon(window_title, icon_path):
    """
    Đổi icon taskbar + title bar cho cửa sổ OpenCV (Windows only)
    """
    try:
        hwnd = win32gui.FindWindow(None, window_title)
        if not hwnd:
            return

        hicon = win32gui.LoadImage(
            None,
            icon_path,
            win32con.IMAGE_ICON,
            0,
            0,
            win32con.LR_LOADFROMFILE | win32con.LR_DEFAULTSIZE
        )

        # ICON nhỏ (title bar)
        win32gui.SendMessage(hwnd, win32con.WM_SETICON, win32con.ICON_SMALL, hicon)
        # ICON lớn (taskbar)
        win32gui.SendMessage(hwnd, win32con.WM_SETICON, win32con.ICON_BIG, hicon)

    except Exception as e:
        print("[ICON ERROR]", e)

def open_aismartmonitor():
    import os, sys, subprocess

    if getattr(sys, 'frozen', False):
        base_dir = os.path.dirname(sys.executable)
    else:
        base_dir = os.path.dirname(os.path.abspath(__file__))

    exe_path = os.path.join(
        base_dir,
        "AISmartMonitor",
        "AISmartMonitor.exe"
    )

    if not os.path.exists(exe_path):
        raise FileNotFoundError("Không tìm thấy AISmartMonitor.exe")

    subprocess.Popen([exe_path], cwd=os.path.dirname(exe_path))

def ask_student_id(parent):
    global ZONE_ID
    import tkinter as tk
    from tkinter import ttk

    result = {"ok": False}

    win = tk.Toplevel(parent)
    win.title("Nhập mã học sinh")
    win.resizable(False, False)
    win.attributes("-topmost", True)

    # ===== CĂN GIỮA =====
    w, h = 360, 180
    x = (win.winfo_screenwidth() - w) // 2
    y = (win.winfo_screenheight() - h) // 2
    win.geometry(f"{w}x{h}+{x}+{y}")

    ttk.Label(
        win,
        text="Nhập Student ID (Zone ID)",
        font=("Segoe UI", 11, "bold")
    ).pack(pady=(15, 5))

    entry = ttk.Entry(win, width=30, justify="center")
    entry.pack(pady=5)
    entry.focus()

    status_label = ttk.Label(win, text="", foreground="red")
    status_label.pack()

    btn_frame = ttk.Frame(win)
    btn_frame.pack(pady=15)

    def confirm(event=None):
        global ZONE_ID
        value = entry.get().strip()

        if not value:
            status_label.config(text="⚠️ Vui lòng nhập Student ID!")
            win.bell()
            return

        ZONE_ID = value
        result["ok"] = True
        win.destroy()

    def cancel():
        win.destroy()

    ttk.Button(btn_frame, text="Xác nhận", command=confirm).pack(side="left", padx=10)
    ttk.Button(btn_frame, text="Hủy", command=cancel).pack(side="left", padx=10)

    # ===== ENTER = XÁC NHẬN =====
    entry.bind("<Return>", confirm)

    win.transient(parent)
    win.grab_set()
    parent.wait_window(win)

    return result["ok"]

def send_incident(state, start_time_str, duration):
    global ZONE_ID

    try:
        now = datetime.datetime.now()
        end_str = now.strftime('%H:%M:%S')
        date_str = now.strftime('%Y-%m-%d')

        payload = {
            "class_id": class_name,
            "zone_id": str(ZONE_ID),
            "issue_type": state,          # Emotion/Posture
            "start_time": start_time_str,
            "end_time": end_str,
            "duration_seconds": duration,
            "date": date_str,
            "scan_mode": SCAN_MODE
        }

        requests.post(SERVER_URL, json=payload, timeout=2)
        print("[INCIDENT SENT]", payload)

    except Exception as e:
        print("[ERROR] Send incident:", e)

def mouse_draw_roi(event, x, y, flags, param):
    global roi_start, roi_end, ROI_DRAWING, ROI_BOX
    global DISPLAY_SCALE_X, DISPLAY_SCALE_Y

    if ROI_ACTIVE:
        return

    if not ROI_DRAWING:
        return
    
    # CHUYỂN TỌA ĐỘ CHUỘT → FRAME GỐC
    fx = int(x * DISPLAY_SCALE_X)
    fy = int(y * DISPLAY_SCALE_Y)

    if event == cv2.EVENT_LBUTTONDOWN and ROI_DRAWING:
        roi_start = (fx, fy)
        roi_end = (fx, fy)

    elif event == cv2.EVENT_MOUSEMOVE and ROI_DRAWING and roi_start:
        roi_end = (fx, fy)

    elif event == cv2.EVENT_LBUTTONUP and ROI_DRAWING:
        roi_end = (fx, fy)

        x1, y1 = roi_start
        x2, y2 = roi_end

        ROI_BOX = (
            min(x1, x2),
            min(y1, y2),
            max(x1, x2),
            max(y1, y2)
        )

        roi_start = None
        roi_end = None

def mouse_draw_roi_fullscreen(event, x, y, flags, param):
    if ROI_ACTIVE:
        return

    global roi_start, roi_end, ROI_BOX, ROI_DRAWING
    global scale_factor

    if not ROI_DRAWING:
        return

    # QUY ĐỔI TỌA ĐỘ TỪ ẢNH HIỂN THỊ → FRAME GỐC
    fx = int(x / scale_factor)
    fy = int(y / scale_factor)

    if event == cv2.EVENT_LBUTTONDOWN:
        roi_start = (fx, fy)
        roi_end = None

    elif event == cv2.EVENT_MOUSEMOVE and roi_start is not None:
        roi_end = (fx, fy)

    elif event == cv2.EVENT_LBUTTONUP and roi_start is not None:
        roi_end = (fx, fy)
        x1, y1 = roi_start
        x2, y2 = roi_end

        ROI_BOX = (
            min(x1, x2),
            min(y1, y2),
            max(x1, x2),
            max(y1, y2)
        )
        roi_start = None
        roi_end = None

def show_success_with_open_folder(parent, export_file_path_csv, summary_path_txt):
    import tkinter as tk
    from tkinter import ttk
    import os, sys, subprocess
    import winsound

    winsound.MessageBeep(winsound.MB_ICONEXCLAMATION)

    folder_path = os.path.dirname(os.path.abspath(export_file_path_csv))

    csv_name = os.path.basename(export_file_path_csv)
    txt_name = os.path.basename(summary_path_txt)

    win = tk.Toplevel(parent)
    win.title("Thành công!")
    win.attributes("-topmost", True)
    win.resizable(False, False)

    # Kích thước hộp thoại
    w, h = 560, 260
    x = (win.winfo_screenwidth() - w) // 2
    y = (win.winfo_screenheight() - h) // 2
    win.geometry(f"{w}x{h}+{x}+{y}")

    frame = ttk.Frame(win)
    frame.pack(fill="both", expand=True, padx=20, pady=15)

    msg = (
        "✅ Đã xuất báo cáo thành công!\n\n"
        f"• File chi tiết (CSV): {csv_name}\n"
        f"• File tổng hợp (TXT): {txt_name}\n\n"
        "📌 Dữ liệu phiên quét hiện tại đã được XÓA\n"
        "để chuẩn bị cho phiên mới."
    )

    label = ttk.Label(
        frame,
        text=msg,
        justify="left",
        wraplength=520
    )
    label.pack(anchor="w")

    # ====== NÚT ======
    btn_frame = ttk.Frame(frame)
    btn_frame.pack(pady=15)

    def open_folder():
        if sys.platform.startswith("win"):
            os.startfile(folder_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", folder_path])
        else:
            subprocess.Popen(["xdg-open", folder_path])

    ttk.Button(btn_frame, text="📂 Mở thư mục", command=open_folder).pack(side="left", padx=12)
    ttk.Button(btn_frame, text="OK", command=win.destroy).pack(side="left", padx=12)

    win.focus_force()

def show_export_success_word(parent, folder_path):
    import tkinter as tk
    from tkinter import ttk
    import os, sys, subprocess, winsound

    winsound.MessageBeep(winsound.MB_ICONASTERISK)

    win = tk.Toplevel(parent)
    win.title("Xuất báo cáo thành công")
    win.resizable(False, False)
    win.attributes("-topmost", True)

    # Căn giữa
    w, h = 520, 200
    x = (win.winfo_screenwidth() - w) // 2
    y = (win.winfo_screenheight() - h) // 2
    win.geometry(f"{w}x{h}+{x}+{y}")

    msg = (
        "✅ Đã xuất báo cáo ROI thành công!\n\n"
        "📁 Thư mục lưu báo cáo:\n"
        f"{folder_path}"
    )

    label = ttk.Label(win, text=msg, wraplength=480, justify="left")
    label.pack(padx=15, pady=15)

    btn_frame = ttk.Frame(win)
    btn_frame.pack(pady=10)

    def open_folder():
        if sys.platform.startswith("win"):
            os.startfile(folder_path)
        elif sys.platform == "darwin":
            subprocess.Popen(["open", folder_path])
        else:
            subprocess.Popen(["xdg-open", folder_path])

    ttk.Button(btn_frame, text="📂 Mở thư mục", command=open_folder).pack(side="left", padx=10)
    ttk.Button(btn_frame, text="OK", command=win.destroy).pack(side="left", padx=10)

    win.focus_force()

def export_roi_to_word():
    global ROI_LOGS, ROI_BOX, log_directory, ROI_IMAGE_BUFFER, ZONE_ID
    from io import BytesIO
    from docx import Document
    from docx.shared import Inches
    from collections import Counter
    import datetime, os
    from PIL import Image
    import matplotlib.pyplot as plt

    if not ROI_LOGS or not ROI_BOX:
        return

    os.makedirs(log_directory, exist_ok=True)

    doc = Document()

    # ===== TIÊU ĐỀ =====
    doc.add_heading(
        f"BÁO CÁO PHÂN TÍCH ROI - EMOTION & POSTURE - HS-{ZONE_ID}", level=1
    )

    # ===== THÔNG TIN CHUNG =====
    now = datetime.datetime.now()
    doc.add_paragraph(f"Thời gian xuất báo cáo: {now.strftime('%d/%m/%Y %H:%M:%S')}")

    start_time = ROI_LOGS[0]['time']
    end_time = ROI_LOGS[-1]['time']
    total_duration = int(end_time - start_time)
    total_frames = len(ROI_LOGS)

    doc.add_paragraph(f"Tổng thời gian quét ROI: {total_duration} giây (~{total_duration/60:.2f} phút)")
    doc.add_paragraph(f"Tổng số frame ghi nhận: {total_frames} frame")

    # ===== HÌNH ẢNH ROI (GIỮ NGUYÊN) =====
    if ROI_IMAGE_BUFFER:
        doc.add_heading("Hình ảnh vùng ROI", level=2)

        image_stream = BytesIO(ROI_IMAGE_BUFFER)
        img = Image.open(image_stream)

        max_width_inch = 6
        max_height_inch = 8
        dpi = 96

        max_width_px = int(max_width_inch * dpi)
        max_height_px = int(max_height_inch * dpi)

        img.thumbnail((max_width_px, max_height_px), Image.Resampling.LANCZOS)

        img_stream_resized = BytesIO()
        img.save(img_stream_resized, format="PNG")
        img_stream_resized.seek(0)

        doc.add_picture(
            img_stream_resized,
            width=Inches(img.width / dpi),
            height=Inches(img.height / dpi)
        )

    # ================== THỐNG KÊ & PHÂN TÍCH ==================

    emo_counter = Counter([log['emotion'] for log in ROI_LOGS])
    posture_list = [log.get('posture') for log in ROI_LOGS if log.get('posture')]
    pos_counter = Counter(posture_list)

    # ===== THỐNG KÊ CẢM XÚC (%) =====
    doc.add_heading("Thống kê biểu cảm khuôn mặt (%)", level=2)

    emo_percent = {}
    for emo, count in emo_counter.items():
        pct = round(count / total_frames * 100, 2)
        emo_percent[emo] = pct
        doc.add_paragraph(f"- {emo}: {pct}%")

    # ===== BIỂU ĐỒ CẢM XÚC =====
    if emo_percent:
        plt.figure(figsize=(6, 4))
        plt.bar(emo_percent.keys(), emo_percent.values())
        plt.title("Phân bố biểu cảm khuôn mặt")
        plt.ylabel("Tỷ lệ (%)")
        plt.xticks(rotation=30)
        plt.tight_layout()

        emo_chart = BytesIO()
        plt.savefig(emo_chart, format="PNG")
        plt.close()
        emo_chart.seek(0)

        doc.add_picture(emo_chart, width=Inches(5))

    # ===== THỐNG KÊ TƯ THẾ (%) =====
    if pos_counter:
        doc.add_heading("Thống kê tư thế (%)", level=2)

        pos_percent = {}
        for pos, count in pos_counter.items():
            pct = round(count / total_frames * 100, 2)
            pos_percent[pos] = pct
            doc.add_paragraph(f"- {pos}: {pct}%")

        # ===== BIỂU ĐỒ TƯ THẾ =====
        plt.figure(figsize=(6, 4))
        plt.bar(pos_percent.keys(), pos_percent.values())
        plt.title("Phân bố tư thế")
        plt.ylabel("Tỷ lệ (%)")
        plt.xticks(rotation=30)
        plt.tight_layout()

        pos_chart = BytesIO()
        plt.savefig(pos_chart, format="PNG")
        plt.close()
        pos_chart.seek(0)

        doc.add_picture(pos_chart, width=Inches(5))

    # ================== TƯ VẤN SỨC KHỎE HỌC ĐƯỜNG ==================

    NEGATIVE_EMOTIONS = ['Buồn', 'Giận dữ', 'Sợ hãi', 'Ghê sợ']

    negative_emo_ratio = sum(
        emo_percent.get(emo, 0) for emo in NEGATIVE_EMOTIONS
    )

    bad_posture_ratio = pos_percent.get('Cúi nhiều (Bad)', 0)
    posture_coverage = sum(pos_percent.values())

    summary_signal_emo = "XANH"
    summary_signal_pos = "XANH"
    quality_check = "TỐT"

    if negative_emo_ratio >= 40:
        summary_signal_emo = "VÀNG (Bất lợi ≥ 40%)"

    if posture_coverage < 50:
        quality_check = "CẦN CẢI THIỆN"
        summary_signal_pos = "VÀNG (Bao phủ < 50%)"
    elif bad_posture_ratio >= 5:
        summary_signal_pos = "VÀNG (Cúi nhiều ≥ 5%)"

    overall_signal = "XANH 🟢"
    if "VÀNG" in summary_signal_emo or "VÀNG" in summary_signal_pos:
        overall_signal = "VÀNG 🟡"
        if "VÀNG" in summary_signal_emo and "VÀNG" in summary_signal_pos:
            overall_signal = "ĐỎ (Nguy cơ kép) 🔴"

    doc.add_heading("Tín hiệu cảnh báo tổng hợp", level=2)
    doc.add_paragraph(f"- Tín hiệu biểu cảm khuôn mặt: {summary_signal_emo}")
    doc.add_paragraph(f"- Tín hiệu tư thế: {summary_signal_pos}")
    doc.add_paragraph(f"- Chất lượng dữ liệu tư thế: {quality_check}")

    doc.add_heading("Đánh giá & tư vấn sức khỏe học đường", level=2)


    doc.add_paragraph(f"Mức độ nguy cơ tổng hợp (phiên quét): {overall_signal}")

    if overall_signal == "XANH 🟢":
        doc.add_paragraph(
            "TỔNG HỢP: Ngưỡng an toàn.\n"
            "KHUYẾN NGHỊ: Duy trì theo dõi định kỳ. "
            "Giáo viên có thể nhắc nhở điều chỉnh tư thế hoặc thay đổi hoạt động nhẹ khi cần."
        )

    elif overall_signal == "VÀNG 🟡":
        doc.add_paragraph(
            "TỔNG HỢP: Nguy cơ trung bình, cần sàng lọc nhanh.\n"
            "QUY TRÌNH ĐỀ XUẤT:\n"
            "• Quan sát bổ sung trong các buổi học tiếp theo.\n"
            "• Nhắc nhở điều chỉnh tư thế, thay đổi hoạt động.\n"
            "• Trao đổi nhẹ nhàng nhằm giảm căng thẳng tâm lý."
        )

    elif overall_signal == "ĐỎ (Nguy cơ kép) 🔴":
        doc.add_paragraph(
            "TỔNG HỢP: Nguy cơ cao, cần kích hoạt tư vấn cá nhân.\n"
            "ĐỀ XUẤT:\n"
            "• Kiểm chứng dữ liệu kỹ thuật và quan sát trực tiếp.\n"
            "• Tham vấn giáo viên chủ nhiệm và chuyên viên tâm lý.\n"
            "• Xây dựng kế hoạch hỗ trợ cá nhân hóa cho học sinh."
        )

    # ===== LƯU FILE =====
    filename = f"ROI_Report_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join(log_directory, filename)
    doc.save(filepath)

    print(f"[INFO] Đã lưu báo cáo ROI tại: {filepath}")
    show_export_success_word(root, log_directory)

def analyze_and_export_csv():
    """
    Phân tích dữ liệu, mở hộp thoại chọn nơi lưu, xuất file CSV và TXT cùng vị trí, 
    khắc phục lỗi font CSV, và reset dữ liệu.
    """
    global DATA_LOGS, history, session_start_time, total_detection_frames, bad_posture_total_frames, BASE_DIR
    from datetime import datetime
    # 0. Kiểm tra dữ liệu hợp lệ (Giữ nguyên kiểm tra)
    if len(DATA_LOGS) < 2 or total_detection_frames == 0:
        return messagebox.showwarning("Thông báo", "Không có đủ dữ liệu để phân tích (cần >1 bản ghi hoặc Frames > 0).")

    # Xác định thời gian phiên
    session_end_time = DATA_LOGS[-1]['timestamp']
    session_duration = session_end_time - session_start_time
    
    # --- A. XỬ LÝ ĐƯỜNG DẪN LƯU FILE (Thực hiện hỏi chỗ lưu) ---
    now = datetime.now()
    date_time_string = now.strftime("%H:%M:%S - Ngày %d/%m/%Y")
    timestamp_str = now.strftime("%Y%m%d_%H%M%S") 
    file_name_base_csv = f"Emotion_Posture_Report_{timestamp_str}.csv"
    file_name_base = f"Emotion_Posture_Report_{timestamp_str}"
    
    # Gán đường dẫn mặc định hoặc đường dẫn người đã thay đổi để lưu file
    export_file_path_csv = os.path.join(log_directory, file_name_base_csv)
        
    # Tạo đường dẫn cho file BÁO CÁO TỔNG HỢP (.txt) cùng thư mục với CSV
    base_dir_csv = os.path.dirname(export_file_path_csv)
    summary_path_txt = os.path.join(base_dir_csv, f"{file_name_base}_SUMMARY.txt")
    
    
    # --- B. PHÂN TÍCH VÀ TÍNH TOÁN ---
    # *Đây là phần phức tạp nhất, tôi giữ lại logic phân tích tổng hợp từ các lần trước để tạo ra file TXT*
    
    emotion_duration = {}
    posture_duration = {}
    for i in range(len(DATA_LOGS) - 1):
        current = DATA_LOGS[i]
        next_record = DATA_LOGS[i+1]
        duration = next_record['timestamp'] - current['timestamp']
        emotion = current.get('emotion', 'Unknown')
        emotion_duration[emotion] = emotion_duration.get(emotion, 0) + duration
        posture = current.get('posture_status', 'N/A')
        posture_duration[posture] = posture_duration.get(posture, 0) + duration

    # Tính Tỷ lệ
    total_valid_emo_duration = sum(dur for emo, dur in emotion_duration.items() if emo not in ['Unknown', 'N/A'])
    total_posture_duration = sum(posture_duration.values())
    emo_ratios = {emo: (dur / total_valid_emo_duration) * 100 for emo, dur in emotion_duration.items()} if total_valid_emo_duration > 0 else {}
    pos_ratios = {pos: (dur / total_posture_duration) * 100 for pos, dur in posture_duration.items()} if total_posture_duration > 0 else {}
    
    NEGATIVE_EMOTIONS = ['Buồn', 'Giận dữ', 'Sợ hãi', 'Ghê sợ'] 
    negative_emo_ratio = sum(emo_ratios.get(emo, 0) for emo in NEGATIVE_EMOTIONS)
    no_posture_duration = posture_duration.get('N/A', 0) + posture_duration.get('Không phát hiện tư thế', 0)
    no_posture_ratio = (no_posture_duration / total_posture_duration) * 100 if total_posture_duration > 0 else 0
    bad_posture_ratio = pos_ratios.get('Cúi nhiều (Bad)', 0)
    posture_coverage = 100 - no_posture_ratio
    fps = total_detection_frames / session_duration if session_duration > 0 else 0
    
    # Quy đổi tín hiệu
    summary_signal_emo = 'XANH'
    summary_signal_pos = 'XANH'
    quality_check = "TỐT"
    if negative_emo_ratio >= 40: summary_signal_emo = 'VÀNG (Bất lợi >= 40%)'
    if posture_coverage < 50:
        quality_check = "CẦN CẢI THIỆN"
        summary_signal_pos = 'VÀNG (Bao phủ < 50%)'
    elif bad_posture_ratio >= 5: summary_signal_pos = 'VÀNG (Cúi nhiều >= 5%)'
    
    # --- LOGIC XÁC ĐỊNH NGƯỠNG TƯ VẤN (TỔNG HỢP - BỔ SUNG) ---
    overall_signal = 'XANH 🟢'
    if 'VÀNG' in summary_signal_emo or 'VÀNG' in summary_signal_pos:
        overall_signal = 'VÀNG 🟡'
        # Trường hợp rủi ro kép (cả hai kênh đều VÀNG), xem như ĐỎ (kích hoạt tư vấn cá nhân) trong bối cảnh báo cáo 1 lần
        if 'VÀNG' in summary_signal_emo and 'VÀNG' in summary_signal_pos:
            overall_signal = 'ĐỎ (Nguy cơ kép) 🔴'

    consultation_recommendation = ""

    if overall_signal == 'XANH 🟢':
        consultation_recommendation = """
    TỔNG HỢP: Ngưỡng an toàn.
    KHUYẾN NGHỊ: Duy trì theo dõi định kỳ.
    Giáo viên có thể nhắc nhở điều chỉnh tư thế hoặc thay đổi hoạt động nhẹ khi cần.
    """
    elif overall_signal == 'VÀNG 🟡':
        consultation_recommendation = """
    TỔNG HỢP: Nguy cơ trung bình, cần sàng lọc nhanh.
    QUY TRÌNH ĐỀ XUẤT:
    • Quan sát bổ sung trong các buổi học tiếp theo.
    • Nhắc nhở điều chỉnh tư thế, thay đổi hoạt động.
    • Trao đổi nhẹ nhàng nhằm giảm căng thẳng tâm lý.
    """
    elif overall_signal == 'ĐỎ (Nguy cơ kép) 🔴':
        consultation_recommendation = """
    TỔNG HỢP: Nguy cơ cao, cần kích hoạt tư vấn cá nhân.
    ĐỀ XUẤT:
    • Kiểm chứng dữ liệu kỹ thuật và quan sát trực tiếp.
    • Tham vấn giáo viên chủ nhiệm và chuyên viên tâm lý.
    • Xây dựng kế hoạch hỗ trợ cá nhân hóa cho học sinh.
    """

    # --- C. TẠO NỘI DUNG BÁO CÁO TỔNG HỢP (TXT) ---
    report_content = f"""
===================================================
| BÁO CÁO TỔNG HỢP PHÂN TÍCH (Mẫu 01a/01b)
| THỜI ĐIỂM XUẤT: {date_time_string}
===================================================
1. THÔNG TIN CHUNG
- Tổng thời gian quét: {session_duration:.2f} giây (~{session_duration/60:.2f} phút)
- Tổng Frames quét: {total_detection_frames}
- Tốc độ khung hình (FPS): {fps:.2f} FPS

2. PHÂN TÍCH TƯ THẾ (POSTURE)
- Tỷ lệ Bao phủ (Posture Coverage): {posture_coverage:.2f}% (Mục tiêu: > 60%)
- Chất lượng dữ liệu Tư thế: {quality_check}
--- Tỷ lệ Chi tiết ---
- Không phát hiện tư thế: {no_posture_ratio:.2f}%
- Ngồi thẳng (Tốt): {pos_ratios.get('Ngồi thẳng (Good)', 0):.2f}%
- Hơi cúi (Cảnh báo): {pos_ratios.get('Hơi cúi (Warning)', 0):.2f}%
- **Cúi nhiều (Bad): {bad_posture_ratio:.2f}%** (Ngưỡng Mẫu 01b: >= 5%)

3. PHÂN TÍCH BIỂU CẢM KHUÔN MẶT
- Tỷ lệ Biểu cảm khuôn mặt Bất lợi: **{negative_emo_ratio:.2f}%** (Ngưỡng Mẫu 01b: >= 40%)
--- Tỷ lệ Chi tiết ---
- Buồn: {emo_ratios.get('Buồn', 0):.2f}%
- Giận dữ: {emo_ratios.get('Giận dữ', 0):.2f}%
- Sợ hãi: {emo_ratios.get('Sợ hãi', 0):.2f}%
- Vui vẻ: {emo_ratios.get('Vui vẻ', 0):.2f}%
- Trung lập: {emo_ratios.get('Trung lập', 0):.2f}%
- Bất ngờ: {emo_ratios.get('Bất ngờ', 0):.2f}%
- Ghê sợ: {emo_ratios.get('Ghê sợ', 0):.2f}%

4. TÍN HIỆU CẢNH BÁO TỔNG HỢP LỚP (Mẫu 01a)
- Tín hiệu Biểu cảm khuôn mặt: {summary_signal_emo}
- Tín hiệu Tư thế: {summary_signal_pos}

===================================================
5. TƯ VẤN VÀ KHUYẾN NGHỊ DỰA TRÊN NGƯỠNG
- Mức độ Nguy cơ Tổng hợp (Phiên quét): {overall_signal}
{consultation_recommendation}
===================================================
"""
    
    # --- D. TẠO NỘI DUNG CSV (Chi tiết) ---
    data_to_export = [
        ["Tổng thời gian quét", f"{session_duration:.2f} giây"],
        ["---", "---"],
        ["Phân tích Biểu cảm", "Thời gian (giây)"],
    ]
    for emo, dur in emotion_duration.items():
        data_to_export.append([emo, f"{dur:.2f}"])
        
    data_to_export.extend([
        ["---", "---"],
        ["Phân tích Tư thế (Tổng hợp)", "Thời gian (giây)"],
    ])
    for pos, dur in posture_duration.items():
        data_to_export.append([pos, f"{dur:.2f}"])
        
    # 5. XUẤT FILE (CSV và TXT)
    try:
        # Xuất file CSV: Dùng encoding='utf-8-sig' để SỬA LỖI CHỮ TIẾNG VIỆT
        with open(export_file_path_csv, 'w', newline='', encoding='utf-8-sig') as csvfile:
            writer = csv.writer(csvfile) 
            writer.writerows(data_to_export) 
        
        # Xuất file BÁO CÁO TỔNG HỢP (.txt)
        with open(summary_path_txt, 'w', encoding='utf-8') as f:
            f.write(report_content)

        show_success_with_open_folder(
            root,# Cửa sổ chính Tk
            export_file_path_csv,
            summary_path_txt
        )


    except Exception as e:
        messagebox.showerror("Lỗi Xuất File", f"Không thể xuất file báo cáo:\n{e}")

    # 6. Xóa dữ liệu và reset biến sau khi xuất
    DATA_LOGS.clear()
    history.clear()
    session_start_time = 0.0
    total_detection_frames = 0
    bad_posture_total_frames = 0

def set_log_directory():
    """Mở hộp thoại để chọn thư mục lưu file log và cập nhật biến toàn cục."""
    global log_directory, root
    
    # Lấy đường dẫn hiện tại làm thư mục ban đầu
    initial_dir = log_directory if os.path.isdir(log_directory) else os.path.expanduser("~")
    
    new_dir = filedialog.askdirectory(
        parent=root,
        initialdir=initial_dir,
        title="Chọn thư mục lưu File Log"
    )
    
    if new_dir:
        log_directory = new_dir
        # Hiển thị thông báo
        messagebox.showinfo("Thành công", f"Thư mục lưu log đã được cập nhật thành công:\n{log_directory}")

def get_local_ip():
    s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
    try:
        s.connect(('8.8.8.8', 80))
        ip = s.getsockname()[0]
    except Exception:
        ip = '127.0.0.1'
    finally:
        s.close()
    return ip

def udp_broadcast(message, port=5000, interval=5):
    """Gửi broadcast UDP liên tục để thiết bị trong LAN bắt được link."""
    sock = socket.socket(socket.AF_INET, socket.SOCK_DGRAM, socket.IPPROTO_UDP)
    sock.setsockopt(socket.SOL_SOCKET, socket.SO_BROADCAST, 1)
    sock.settimeout(0.2)
    while True:
        try:
            sock.sendto(message.encode("utf-8"), ("<broadcast>", port))
        except Exception as e:
            print("Broadcast error:", e)
        time.sleep(interval)

def list_cameras():
    """Liệt kê camera sử dụng pygrabber.dshow_graph."""
    graph = FilterGraph()
    devices = graph.get_input_devices()
    return devices

def draw_text_with_outline(draw, pos, text, font, text_color,
                           outline_color=(0, 0, 0), outline_width=1):
    x, y = pos
    for dx in range(-outline_width, outline_width + 1):
        for dy in range(-outline_width, outline_width + 1):
            if dx == 0 and dy == 0:
                continue
            draw.text((x + dx, y + dy), text, font=font, fill=outline_color)
    draw.text((x, y), text, font=font, fill=text_color)

def draw_filled_rectangle_with_outline(img, pt1, pt2, color,
                                      outline_color=(0, 0, 0),
                                      outline_width=1):
    cv2.rectangle(img,
                  (pt1[0] - outline_width, pt1[1] - outline_width),
                  (pt2[0] + outline_width, pt2[1] + outline_width),
                  outline_color, -1)
    cv2.rectangle(img, pt1, pt2, color, -1)

def calculate_angle(a, b, c):
    import math
    ax, ay = a
    bx, by = b
    cx, cy = c
    angle = math.degrees(
        math.atan2(cy - by, cx - bx) - math.atan2(ay - by, ax - bx)
    )
    return abs(angle)

def bring_window_to_front(window_name):
    """Set always on top + nhảy ra trước."""
    hwnd = win32gui.FindWindow(None, window_name)
    if hwnd:
        win32gui.SetWindowPos(hwnd, win32con.HWND_TOPMOST,
                              0, 0, 0, 0,
                              win32con.SWP_NOMOVE | win32con.SWP_NOSIZE)
        win32gui.ShowWindow(hwnd, win32con.SW_SHOWNORMAL)
        try:
            win32gui.SetForegroundWindow(hwnd)
        except Exception:
            pass

def show_warning(msg):
    """Hiện cảnh báo (luôn hiện trên cùng)."""
    win = tk.Toplevel()
    win.withdraw()
    win.attributes('-topmost', True)
    messagebox.showwarning("Cảnh báo", msg, parent=win)
    win.destroy()

# UTILITIES CHO TKINTER & LOADING

def generate_qr_code(link):
    """Tạo mã QR từ link và trả về dưới dạng đối tượng PIL Image."""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10, # Tăng kích thước box_size
        border=4,
    )
    qr.add_data(link)
    qr.make(fit=True)
    # Trả về ảnh PIL Image, không phải PhotoImage
    img_qr_pil = qr.make_image(fill_color="black", back_color="white").convert('RGB')
    return img_qr_pil

# Copy đường link flask server
def copy_link_to_clipboard(link, link_window):
    global root
    root.clipboard_clear()
    root.clipboard_append(link)
    messagebox.showinfo("Thông báo", "Đã sao chép đường link vào Clipboard!")
    link_window.destroy() 

# Copy mã qr
def copy_qr_to_clipboard(qr_image_pil, link_window):
    """
    Sao chép ảnh PIL Image (QR Code) vào Clipboard dưới dạng DIB (Bitmap).
    CHỈ hoạt động trên Windows vì sử dụng win32clipboard.
    """
    try:
        # Chuyển đổi PIL Image sang định dạng BMP byte stream
        output = io.BytesIO()
        qr_image_pil.save(output, 'BMP')
        data = output.getvalue()[14:] # Bỏ qua BMP file header (14 bytes)
        
        # Mở Clipboard và đặt dữ liệu
        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
        # Đặt format CF_DIB (Device Independent Bitmap)
        win32clipboard.SetClipboardData(win32con.CF_DIB, data)
        win32clipboard.CloseClipboard()
        
        messagebox.showinfo("Thông báo", "Đã sao chép ảnh QR Code vào Clipboard thành công!")
        link_window.destroy()
        
    except Exception as e:
        messagebox.showerror("Lỗi Sao Chép Ảnh", f"Không thể sao chép ảnh QR Code vào Clipboard (Chỉ hỗ trợ Windows).\nLỗi: {e}")

# Hiện đường link
def show_stream_link(link):
    """Hiển thị đường link Stream, Mã QR và các nút hành động."""
    global root
    
    # 1. Tạo ảnh QR dưới dạng PIL Image
    qr_image_pil = generate_qr_code(link)
    
    # 2. Chuyển đổi sang PhotoImage để hiển thị trong Tkinter
    bio = io.BytesIO()
    qr_image_pil.save(bio, format='PNG')
    qr_photo = tk.PhotoImage(data=bio.getvalue())

    # 3. Tạo cửa sổ Toplevel
    link_window = tk.Toplevel(root)
    link_window.title("Đường Link Stream và Mã QR")
    link_window.update() 

    # 4. Hiển thị Mã QR
    qr_label = tk.Label(link_window, image=qr_photo)
    qr_label.image = qr_photo # Giữ tham chiếu để tránh bị Garbage Collection
    qr_label.pack(pady=10, padx=20)
    
    # 5. Hiển thị Text Label và Entry
    tk.Label(link_window, text="Quét Mã QR hoặc truy cập đường link sau:", 
             font=("Arial", 10, "bold")).pack(pady=(0, 5), padx=20)
    
    link_entry = tk.Entry(link_window, width=50, justify='center')
    link_entry.insert(0, link)
    link_entry.config(state="readonly")
    link_entry.pack(pady=5, padx=20)
    
    # 6. Khung chứa các nút
    button_frame = tk.Frame(link_window)
    button_frame.pack(pady=15)
    
    # Nút 1: Sao chép Link
    copy_link_btn = tk.Button(button_frame,
                         text="Sao chép Link",
                         command=lambda: copy_link_to_clipboard(link, link_window),
                         bg="#007ACC", fg="white", font=("Arial", 10, "bold"))
    copy_link_btn.pack(side=tk.LEFT, padx=5)

    # Nút 2: Sao chép Ảnh QR vào Clipboard (Mới)
    copy_qr_btn = tk.Button(button_frame,
                         text="📋 Sao chép Ảnh QR",
                         command=lambda: copy_qr_to_clipboard(qr_image_pil, link_window),
                         bg="#FF9800", fg="white", font=("Arial", 10, "bold"))
    copy_qr_btn.pack(side=tk.LEFT, padx=5)

    # Nút 3: Đóng
    close_btn = tk.Button(button_frame,
                         text="Đóng",
                         command=link_window.destroy,
                         bg="#F44336", fg="white", font=("Arial", 10, "bold"))
    close_btn.pack(side=tk.LEFT, padx=5)
    
    # 7. Căn giữa cửa sổ mới
    root.update_idletasks()
    link_window.update_idletasks()
    x = root.winfo_x() + (root.winfo_width() - link_window.winfo_reqwidth()) // 2
    y = root.winfo_y() + (root.winfo_height() - link_window.winfo_reqheight()) // 2
    link_window.geometry(f"+{x}+{y}")

# Hiện hộp thoại loading
def show_loading_window(title="Đang khởi động hệ thống..."):
    global loading_window, progress_bar, progress_label, root

    loading_window = tk.Toplevel(root)
    loading_window.title(title)
    loading_window.geometry("400x140")
    loading_window.resizable(False, False)
#    loading_window.attributes('-topmost', True)

    tk.Label(loading_window, text="Đang khởi động hệ thống, vui lòng chờ...",
             font=("Arial", 10)).pack(pady=10)

    progress_bar = ttk.Progressbar(loading_window, orient="horizontal", length=350, mode="determinate")
    progress_bar.pack(pady=10)
    progress_bar["maximum"] = 100
    progress_bar["value"] = 0

    progress_label = tk.Label(loading_window, text="0%", font=("Arial", 10, "bold"))
    progress_label.pack()

    # Căn giữa cửa sổ loading
    root.update_idletasks()
    loading_window.update_idletasks()
    x = root.winfo_x() + (root.winfo_width() - loading_window.winfo_reqwidth()) // 2
    y = root.winfo_y() + (root.winfo_height() - loading_window.winfo_reqheight()) // 2
    loading_window.geometry(f"+{x}+{y}")

def update_progress(percent, text=None):
    if progress_bar and progress_label and loading_window and loading_window.winfo_exists():
        progress_bar["value"] = percent
        if text:
            progress_label.config(text=f"{text} ({percent}%)")
        else:
            progress_label.config(text=f"{percent}%")
        loading_window.update_idletasks()

def destroy_loading_window():
    global loading_window
    if loading_window and loading_window.winfo_exists():
        loading_window.destroy()

# FLASK STREAMING SETUP

def start_flask_server():
    global flask_app
    flask_app.run(host='0.0.0.0', port=5000, threaded=True, debug=False, use_reloader=False)

def gen_frames():
    global latest_frame, frame_lock, is_running
    while is_running:
        with frame_lock:
            if latest_frame is None:
                time.sleep(0.03)
                continue
            quality = 70 if current_mode == 'camera' else 50
            ret, buffer = cv2.imencode('.jpg', latest_frame, [int(cv2.IMWRITE_JPEG_QUALITY), quality])
            if not ret:
                continue
            jpg = buffer.tobytes()
        yield (b'--frame\r\n'
               b'Content-Type: image/jpeg\r\n\r\n' + jpg + b'\r\n')
        time.sleep(0.03) 

@flask_app.route('/')
def index():
    icon_name = "Emotion + Posture Detector v3.0 Camera.ico" if current_mode == 'camera' else "Emotion + Posture Detector v3.0 Fullscreen Capture.ico."
    title_text = "Camera" if current_mode == 'camera' else "Fullscreen Capture"
    html_page = HTML_PAGE.replace('{{ title_type }}', title_text).replace('{{ icon_name }}', icon_name)
    return render_template_string(html_page)

@flask_app.route('/video_feed')
def video_feed():
    return Response(gen_frames(), mimetype='multipart/x-mixed-replace; boundary=frame')

# Trang web (Chung)
HTML_PAGE = """
<html>
  <head>
    <link rel="icon" href="{{ url_for('static', filename='{{ icon_name }}') }}" type="image/x-icon">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>Emotion + Posture Detector Stream</title>
    <style>
      /* Đảm bảo body chiếm toàn bộ viewport */
      html, body {
        height: 100%; /* Chiều cao 100% của viewport */
        width: 100%;  /* Chiều rộng 100% của viewport */
        margin: 0;
        padding: 0;
        overflow: hidden; /* Ngăn cuộn trang nếu ảnh quá lớn */
      }

      body {
        background: #a19fa2;
        color: #fff;
        font-family: Arial, sans-serif;
        display: flex;
        flex-direction: column;
        justify-content: center; /* Căn giữa theo chiều dọc */
        align-items: center;    /* Căn giữa theo chiều ngang */
      }

      h2 {
        margin-top: 20px; /* Thêm khoảng cách trên cho tiêu đề */
        margin-bottom: 20px;
        flex-shrink: 0; /* Đảm bảo tiêu đề không bị co lại */
      }

      /* Container cho ảnh để nó chiếm không gian còn lại */
      .video-container {
        flex-grow: 1; /* Cho phép container chiếm hết không gian còn lại */
        width: 100%;
        display: flex;
        justify-content: center;
        align-items: center;
        padding: 10px; /* Thêm padding nhẹ xung quanh ảnh */
        box-sizing: border-box;
      }

      img {
        /* Kích thước tối đa là 100% của container chứa nó */
        max-width: 100%;
        max-height: 100%;
        /* Tự động điều chỉnh kích thước để toàn bộ ảnh hiển thị mà không bị cắt */
        object-fit: contain;
        /* Kích thước ảnh thực tế */
        width: auto;
        height: auto;
        /* Giữ lại các style gốc */
        border-radius: 10px;
        box-shadow: 0 0 15px rgba(0, 0, 0, 0.3);
      }
    </style>
  </head>
  <body>
    <h2>Emotion + Posture Detector Live - {{ title_type }}</h2>
    <img src="{{ url_for('video_feed') }}">
  </body>
</html>
"""

# HÀM HỎI CÓ QUÉT TIẾP HAY KHÔNG(DÙNG CHUNG CHO CAMERA VÀ FULLSCREEN)
def ask_yes_no_blocking(title, message):
    dialog = tk.Toplevel(root)
    dialog.title(title)
    dialog.attributes('-topmost', True)
    dialog.grab_set()  # khóa focus
    dialog.resizable(False, False)

    result = {'value': False}

    def on_yes():
        result['value'] = True
        dialog.destroy()

    def on_no():
        result['value'] = False
        dialog.destroy()

    tk.Label(dialog, text=message, justify='left', wraplength=400)\
        .pack(padx=20, pady=15)

    btn_frame = tk.Frame(dialog)
    btn_frame.pack(pady=10)

    tk.Button(btn_frame, text="Yes", width=10, command=on_yes)\
        .pack(side='left', padx=10)
    tk.Button(btn_frame, text="No", width=10, command=on_no)\
        .pack(side='right', padx=10)

    # Căn giữa màn hình
    dialog.update_idletasks()
    w = dialog.winfo_width()
    h = dialog.winfo_height()
    x = (dialog.winfo_screenwidth() // 2) - (w // 2)
    y = (dialog.winfo_screenheight() // 2) - (h // 2)
    dialog.geometry(f"+{x}+{y}")

    dialog.wait_window()  # ⛔ BLOCK tại đây

    return result['value']

def ask_yes_no_on_main_thread(title, message):
    result = {'value': None}

    def _show():
        result['value'] = ask_yes_no_blocking(title, message)

    # Chạy trên MAIN THREAD
    root.after(0, _show)

    # CHỜ kết quả
    while result['value'] is None:
        time.sleep(0.05)

    return result['value']

# HÀM CHÍNH CHO CAMERA

def run_detection_camera(cam_index):
    global latest_frame, frame_lock, is_running, root, broadcast_thread, detection_thread
    global cap, DATA_LOGS, SCAN_MIN_DURATION
    global ROI_ACTIVE, ROI_BOX, ROI_DRAWING, ROI_IMAGE_PATH
    global roi_status_color, roi_status_text, roi_emotion_label, ABNORMAL_THRESHOLD
    global class_name, ZONE_ID
    global roi_start, roi_end
    global DISPLAY_SCALE_X, DISPLAY_SCALE_Y
    global force_exit_no_report

    ROI_ACTIVE = False
    ROI_DRAWING = False
    DATA_LOGS = [] # Xóa log cũ
    scan_start_time = time.time()
    last_log_time = time.time()

    window_title = 'Emotion + Posture Detector v5.0 (Camera)'
    
    # KHI HÀM BẮT ĐẦU CHẠY: Báo cho GUI biết là detection đang chạy
    with thread_lock:
        is_running = True

    local_ip = get_local_ip()
    link = f"http://{local_ip}:5000/"
    
    if not hasattr(run_detection_camera, "_flask_started"):
        Thread(target=start_flask_server, daemon=True).start()
        time.sleep(1)
        run_detection_camera._flask_started = True

    if broadcast_thread is None or not broadcast_thread.is_alive():
        broadcast_thread = Thread(target=udp_broadcast, args=(link,), daemon=True)
        broadcast_thread.start()

    update_progress(25, "Đang tải mô hình nhận diện biểu cảm (Keras)...")
    import mediapipe as mp
    try:
        from tensorflow.keras.models import load_model
    except ImportError:
        messagebox.showerror("Lỗi", "Vui lòng cài đặt Tensorflow/Keras.")
        root.after(0, destroy_loading_window)
        with thread_lock:
             is_running = False
        return

    face_xml = os.path.join(BASE_DIR, "haarcascade_frontalface_default.xml")
    model_h5 = os.path.join(BASE_DIR, "emotion_detection.h5")
    face_classifier = cv2.CascadeClassifier(face_xml)
    if face_classifier.empty():
        messagebox.showerror("Lỗi", f"Không tìm thấy file cascade: {face_xml}")
        root.after(0, destroy_loading_window)
        with thread_lock:
             is_running = False
        return

    classifier = load_model(model_h5)
    update_progress(50, "Đang tải mô hình tư thế (MediaPipe)...")
    class_labels = ['Giận dữ', 'Ghê sợ', 'Sợ hãi', 'Vui vẻ', 'Buồn', 'Bất ngờ', 'Trung lập']
    mp_pose = mp.solutions.pose
    pose = mp_pose.Pose(min_detection_confidence=0.5, min_tracking_confidence=0.5)
    mp_drawing = mp.solutions.drawing_utils

    update_progress(70, "Đang mở camera...")
    cap = cv2.VideoCapture(cam_index, cv2.CAP_DSHOW)
    WIDTH, HEIGHT = 1280, 720
    cap.set(cv2.CAP_PROP_FRAME_WIDTH, WIDTH)
    cap.set(cv2.CAP_PROP_FRAME_HEIGHT, HEIGHT)
    if not cap.isOpened():
        messagebox.showerror("Lỗi", "Không thể mở camera.")
        root.after(0, destroy_loading_window)
        with thread_lock:
             is_running = False
        return

    update_progress(100, "Hoàn tất! Mở camera...")
    root.after(0, destroy_loading_window)
    root.after(100, show_stream_link, link)
    
    global session_start_time, bad_posture_total_frames, total_detection_frames, history
    session_start_time = time.time()
    bad_posture_total_frames = 0
    total_detection_frames = 0
    history.clear() # Đảm bảo history sạch khi bắt đầu phiên mới

    start_time = time.time()
    interval = 120
    scale_factor = 1.0
    first_show = True

    font = ImageFont.truetype(font_path, 28)
    font2 = ImageFont.truetype(font_path, 20)

    short_term_emotion_buffer = deque(maxlen=STABILITY_WINDOW_FRAMES)
    bad_posture_counter = 0 
    current_stable_emotion = 'Trung lập'

    status_posture = "Không phát hiện tư thế"

    force_exit_no_report = False
    # Vòng lặp chính
    while cap.isOpened() and is_running:
        ret, frame = cap.read()
        if not ret: break

        current_time = time.time()

        # ... (Phần logic Emotion, Posture Detection và Drawing giữ nguyên) ...
        # --- Emotion Detection ---
        frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(frame_pil)
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        faces = face_classifier.detectMultiScale(gray, 1.3, 5)

        final_emotion_label = current_stable_emotion
        roi_emotion_label = current_stable_emotion

        for (x, y, w, h) in faces:
            face_cx = x + w // 2
            face_cy = y + h // 2

            face_in_roi = False
            if ROI_ACTIVE and ROI_BOX:
                rx1, ry1, rx2, ry2 = ROI_BOX
                if rx1 <= face_cx <= rx2 and ry1 <= face_cy <= ry2:
                    face_in_roi = True

            cv2.rectangle(frame, (x, y), (x + w, y + h), (242, 248, 68), 2)
            roi_gray = cv2.resize(gray[y:y + h, x:x + w], (48, 48), interpolation=cv2.INTER_AREA)
            roi = np.expand_dims(np.expand_dims(roi_gray.astype("float") / 255.0, axis=-1), axis=0)
            
            # Lấy xác suất
            predictions = classifier.predict(roi, verbose=0)
            probabilities = predictions[0] 
            sorted_indices = np.argsort(probabilities)[::-1]
            
            p_max = probabilities[sorted_indices[0]]
            p_top2 = probabilities[sorted_indices[1]]
            predicted_label = class_labels[sorted_indices[0]] # Nhãn dự đoán

            # 1. KIỂM TRA NGƯỠNG (VALIDATION) theo DOCX
            config = EMOTION_THRESHOLDS.get(predicted_label, {'p_max': THRESHOLD_P_MAX_DEFAULT, 'delta': THRESHOLD_DELTA_TOP2_DEFAULT})
            
            is_reliable = p_max >= config['p_max']
            is_not_ambiguous = (p_max - p_top2) >= config['delta']
            
            if is_reliable and is_not_ambiguous:
                raw_validated_label = predicted_label
            else:
                raw_validated_label = 'Unknown' # Nhãn không đủ tin cậy (theo DOCX)

            # SỬA CHỮA ĐỂ ĐỒNG BỘ HÓA LOGIC ỔN ĐỊNH CẢM XÚC:
            # 2. CỬA SỔ ỔN ĐỊNH NGẮN HẠN (W: 1 FRAME HOẶC ỔN ĐỊNH THEO CẤU HÌNH)
            short_term_emotion_buffer.append(raw_validated_label)
            final_emotion_label = current_stable_emotion # Giữ nhãn cũ trước khi có kết quả mới

            if len(short_term_emotion_buffer) == STABILITY_WINDOW_FRAMES:
                counts = Counter(short_term_emotion_buffer)
                most_common_label, count = counts.most_common(1)[0]
                
                # Nếu nhãn phổ biến nhất chiếm ưu thế (>= 80%) và không phải Unknown
                if most_common_label != 'Unknown' and (count / STABILITY_WINDOW_FRAMES) >= STABILITY_DOMINANCE_RATIO: 
                    final_emotion_label = most_common_label
                # Ngược lại, giữ nhãn ổn định trước đó
                else: 
                    final_emotion_label = current_stable_emotion 
                
                # Cập nhật nhãn ổn định hiện tại và lưu vào lịch sử dài hạn
                current_stable_emotion = final_emotion_label
                
                # Cập nhật lịch sử dài hạn (Chỉ để tính Negative Ratio)
                if current_stable_emotion != 'Unknown':
                     # 1 = Tiêu cực; 0 = Tích cực/Trung lập
                     if current_stable_emotion in ['Giận dữ', 'Ghê sợ', 'Sợ hãi', 'Buồn']:
                         history.append(1)
                     else:
                         history.append(0)
                
                # Xóa buffer để bắt đầu cửa sổ mới
                short_term_emotion_buffer.clear()
            else:
                final_emotion_label = current_stable_emotion # Hiển thị nhãn ổn định trước đó
            
            if face_in_roi:
                roi_emotion_label = final_emotion_label
            # Vẽ nhãn cảm xúc đã được ổn định lên frame
            frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            draw = ImageDraw.Draw(frame_pil)
            draw_text_with_outline(draw, (x, y - 35), final_emotion_label, font, (0, 255, 0))
            frame = cv2.cvtColor(np.array(frame_pil), cv2.COLOR_RGB2BGR)

        negative_ratio = sum(history) / len(history) if len(history) > 0 else 0
        elapsed = time.time() - start_time
        #if elapsed >= interval:
        #    if negative_ratio > 0.6:
        #        show_warning("Pause / đổi hoạt động / nghỉ 2 phút")
        #    start_time = time.time()
        

        labels = [final_emotion_label]
        # --- Posture Detection ---

        if ROI_ACTIVE and ROI_BOX:
            x1, y1, x2, y2 = ROI_BOX

            # Cắt ảnh theo ROI
            roi_frame = frame[y1:y2, x1:x2]

            if roi_frame.size == 0:
                results = None
            else:
                image_rgb = cv2.cvtColor(roi_frame, cv2.COLOR_BGR2RGB)
                results = pose.process(image_rgb)
        else:
            image_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            results = pose.process(image_rgb)   


        angle_back = 0
        status_back_detail = "Không phát hiện tư thế"
        status_posture = "Không phát hiện tư thế"
        color = (255, 255, 255)
        angle_back, angle_neck, angle_elbow = 0, 0, 0
        status_back, status_neck, status_elbow = "Không phát hiện", "Không phát hiện", "Không phát hiện"
        status_posture = "Không phát hiện tư thế"

        if results is not None and results.pose_landmarks:
            landmarks = results.pose_landmarks.landmark
            person_in_roi = False

            if ROI_ACTIVE and ROI_BOX:
                x1, y1, x2, y2 = ROI_BOX
                roi_w = x2 - x1
                roi_h = y2 - y1

                # Landmark đang là tọa độ TRONG ROI
                cx = int(
                    ((landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].x +
                      landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].x) * 0.5) * roi_w
                ) + x1

                cy = int(
                    ((landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].y +
                      landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].y) * 0.5) * roi_h
                ) + y1

                if x1 <= cx <= x2 and y1 <= cy <= y2:
                    person_in_roi = True


            if (not ROI_ACTIVE) or person_in_roi:
                try:
                    ear_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].y
                    ]
                    shoulder_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].y
                    ]
                    hip_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_HIP.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_HIP.value].y
                    ]

                    ear_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].y
                    ]
                    shoulder_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].y
                    ]
                    hip_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value].y
                    ]

                    vis_l = landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].visibility
                    vis_r = landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].visibility

                    angle_back = None

                    if vis_l >= vis_r and vis_l > 0.5:
                        angle_back = calculate_angle(hip_l, shoulder_l, ear_l)
                    elif vis_r > 0.5:
                        angle_back = calculate_angle(hip_r, shoulder_r, ear_r)
                    else:
                        angle_back = None  # Không đủ tin cậy


                    # --- Quy đổi Ergonomics (Dựa trên Góc Lưng) ---
                    if angle_back is None:
                        status_posture = "Không đủ dữ liệu"
                        color = (255, 255, 255)

                    else:
                        if angle_back >= 170:
                            status_back_detail = "Ngồi thẳng (Good)"
                            status_posture = "Ngồi thẳng (Good)"
                            color = (0, 255, 0) # Xanh
                        elif 150 <= angle_back < 170:
                            status_back_detail = "Hơi cúi (Warning)"
                            status_posture = "Hơi cúi (Warning)"
                            color = (255, 255, 0) # Vàng
                        else:
                            status_back_detail = "Cúi nhiều (Bad)"
                            status_posture = "Cúi nhiều (Bad)"
                            color = (255, 0, 0) # Đỏ
                    
                    # Vẽ Landmarks
                    if ROI_ACTIVE and ROI_BOX:
                        mp_drawing.draw_landmarks(
                            roi_frame,
                            results.pose_landmarks,
                            mp_pose.POSE_CONNECTIONS
                        )
                        # dán ROI đã vẽ landmark lại frame gốc
                        frame[y1:y2, x1:x2] = roi_frame
                    else:
                        mp_drawing.draw_landmarks(
                            frame,
                            results.pose_landmarks,
                            mp_pose.POSE_CONNECTIONS
                        )

                except Exception:
                    #Xử lý khi không tìm thấy đủ 3 điểm (Xảy ra khi quay nghiêng quá nhiều)
                    status_back_detail = "Không đủ điểm (Quá nghiêng)"
                    status_posture = "Không phát hiện tư thế"
                    color = (255, 255, 255) # Trắng

            if ROI_ACTIVE and person_in_roi:
                # Cho file word
                ROI_LOGS.append({ 
                        "time": time.time(), 
                        "emotion": roi_emotion_label, 
                        "posture": status_posture 
                    })

                # Cho trang web
                current_state = f"{roi_emotion_label}/{status_posture}"
                now = time.time()

                # Nếu chưa có trạng thái → khởi tạo
                if ROI_STATE_TRACKER["state"] is None:
                    ROI_STATE_TRACKER["state"] = current_state
                    ROI_STATE_TRACKER["start_time"] = now

                # Nếu trạng thái THAY ĐỔI
                elif current_state != ROI_STATE_TRACKER["state"]:
                    prev_state = ROI_STATE_TRACKER["state"]
                    start_time_roi = ROI_STATE_TRACKER["start_time"]
                    duration = int(now - start_time_roi)

                    # Kiểm tra trạng thái cũ có bất thường không
                    is_abnormal = (
                        "Warning" in prev_state or "Bad" in prev_state or
                        any(x in prev_state for x in ["Buồn", "Giận dữ", "Ghê sợ", "Sợ hãi"])
                    )

                    # Nếu bất thường và đủ thời gian → GỬI NGAY
                    if is_abnormal and duration >= ABNORMAL_THRESHOLD:
                        now_dt = datetime.datetime.now()
                        start_time_str = (now_dt - datetime.timedelta(seconds=duration)).strftime('%H:%M:%S')

                        send_incident(
                            prev_state,
                            start_time_str,
                            duration
                        )

                    # Reset sang trạng thái mới
                    ROI_STATE_TRACKER["state"] = current_state
                    ROI_STATE_TRACKER["start_time"] = now


        if "Bad" in status_posture:
            bad_posture_counter += 1
        else:
            bad_posture_counter = 0

        if bad_posture_counter >= BAD_POSTURE_WARNING_FRAMES:
            show_warning("CẢNH BÁO TƯ THẾ: Bạn đã ngồi cúi gù quá lâu. Vui lòng điều chỉnh lại tư thế ngồi!")
            bad_posture_counter = 0

        total_detection_frames += 1 # Tăng tổng frame

        box_color = (0, 0, 255) if negative_ratio > 0.6 else (0, 255, 255) if 0.2 <= negative_ratio <= 0.6 else (0, 255, 0)

        # --- LOGGING DỮ LIỆU ---
        # Ghi log mỗi 0.5 giây để tiết kiệm tài nguyên
        if current_time - last_log_time >= 0.5 and labels:
            with LOG_LOCK:
                # SỬ DỤNG status_posture ĐÃ ĐƯỢC XÁC ĐỊNH
                DATA_LOGS.append({
                    'timestamp': current_time,
                    'emotion': labels[0] if labels else 'Không phát hiện', 
                    'posture_status': status_posture,
                })
            last_log_time = current_time

        frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(frame_pil)

        draw_text_with_outline(draw, (20, 20), f"Tư thế Lưng: {status_back_detail} ({int(angle_back)}°)", font2, (255, 255, 255))
        draw_text_with_outline(draw, (20, 50), f"Trạng thái tổng: {status_posture}", font2, color) # Hiển thị tóm tắt và màu
        draw_text_with_outline(draw, (20, 80), f"Số lượng: {len(faces)}", font2, (255, 0, 255))
        draw_text_with_outline(draw, (20, 110), "Trạng thái:", font2, (0, 0, 255))
        draw_text_with_outline(draw, (960, 680), "Bấm phím 'Q' để thoát", font, (255, 255, 0)) # Chú thích thay đổi
        draw_text_with_outline(draw, (1000, 20), "Bấm phím 'M' để phóng to", font2, (255, 255, 0))
        draw_text_with_outline(draw, (1000, 50), "Bấm phím 'N' để thu nhỏ", font2, (255, 255, 0))
        draw_text_with_outline(draw, (1000, 80), "Đang quét tại: " + class_name, font2, (0, 255, 0))
        draw_text_with_outline(draw, (1000, 110), roi_status_text, font2, roi_status_color)
        frame = cv2.cvtColor(np.array(frame_pil), cv2.COLOR_RGB2BGR)
        draw_filled_rectangle_with_outline(frame, (195 - 65, 125 - 13), (215 - 65, 145 - 13), box_color, outline_width=2)
        
        # ... (Phần xử lý phím bấm và hiển thị giữ nguyên) ...
        key = cv2.waitKey(10) & 0xFF
        if key == ord('m'): scale_factor = min(1.0, scale_factor + 0.1)
        elif key == ord('n'): scale_factor = max(0.2, scale_factor - 0.1)
        elif key == ord('q'):
            current_time = time.time()
            elapsed = current_time - scan_start_time

            # CHƯA đủ thời gian quét
            if elapsed < SCAN_MIN_DURATION:
                remaining = int(SCAN_MIN_DURATION - elapsed)

                answer = ask_yes_no_blocking(
                    "Chưa đủ thời gian quét",
                    f"Cần quét tối thiểu {SCAN_MIN_DURATION} giây.\n"
                    f"Bạn cần quét thêm {remaining} giây nữa.\n\n"
                    f"Nếu dừng bây giờ sẽ KHÔNG xuất báo cáo.\n"
                    f"Bạn có chắc chắn muốn dừng không?"
                )

                if answer:
                    force_exit_no_report = True
                    break   # thoát vòng lặp CV2
                else:
                    continue  # tiếp tục quét

            # ĐÃ đủ thời gian quét
            else:
                break

        elif key == ord('v') and not ROI_ACTIVE:
            ROI_DRAWING = not ROI_DRAWING

            # CHỈ xoá khung khi TẮT vẽ
            if not ROI_DRAWING:
                roi_start = None
                roi_end = None
                ROI_BOX = None

        elif key == ord('s') and ROI_BOX and not ROI_ACTIVE:
            global ROI_IMAGE_BUFFER
            if not ZONE_ID:
                ok = ask_student_id(root)

                # ÉP cửa sổ OpenCV hiện lại sau khi đóng Tkinter
                cv2.namedWindow(window_title)
                hwnd = win32gui.FindWindow(None, window_title)
                if hwnd:
                    win32gui.SetWindowPos(
                        hwnd,
                        win32con.HWND_TOPMOST,
                        0, 0, 0, 0,
                        win32con.SWP_NOMOVE | win32con.SWP_NOSIZE
                    )

                if not ok:
                    show_warning("Vui lòng nhập Student ID trước khi quét ROI.")
                    continue


            ROI_ACTIVE = True
            ROI_LOGS.clear()
            roi_scan_start_time = time.time()

            x1, y1, x2, y2 = ROI_BOX
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(frame.shape[1], x2), min(frame.shape[0], y2)

            roi_crop = frame[y1:y2, x1:x2]

            if roi_crop.size > 0:
                _, buffer = cv2.imencode(".png", roi_crop)
                ROI_IMAGE_BUFFER = buffer.tobytes()

            show_warning("BẮT ĐẦU QUÉT ROI (ĐÃ CHỤP ẢNH ROI)")


        elif key == ord('e') and ROI_ACTIVE:
            # GỬI TRẠNG THÁI CUỐI CÙNG NẾU ĐỦ ĐIỀU KIỆN
            if ROI_STATE_TRACKER["state"] and ROI_STATE_TRACKER["start_time"]:
                now = time.time()
                duration = now - ROI_STATE_TRACKER["start_time"]

                if duration >= ABNORMAL_THRESHOLD:
                    duration = int(time.time() - ROI_STATE_TRACKER["start_time"])

                    now_dt = datetime.datetime.now()
                    start_time_str = (now_dt - datetime.timedelta(seconds=duration)).strftime('%H:%M:%S')

                    send_incident(
                        ROI_STATE_TRACKER["state"],
                        start_time_str,
                        duration
                    )

            ROI_ACTIVE = False
            ROI_STATE_TRACKER["state"] = None
            ROI_STATE_TRACKER["start_time"] = None

            export_roi_to_word()
            ROI_IMAGE_PATH = None
            ZONE_ID = None

        if ROI_DRAWING:
            roi_status_text = "Vẽ khung ROI: BẬT"
            roi_status_color = (0, 255, 0)
        else:
            roi_status_text = "Vẽ khung ROI: TẮT"
            roi_status_color = (255, 0, 0)


        frame_stream = cv2.resize(frame.copy(), (int(frame.shape[1]*0.6), int(frame.shape[0]*0.6)))
        with frame_lock:
            latest_frame = frame_stream.copy()

        new_w, new_h = int(WIDTH * scale_factor), int(HEIGHT * scale_factor)

        DISPLAY_SCALE_X = WIDTH / new_w
        DISPLAY_SCALE_Y = HEIGHT / new_h

        # ===== VẼ ROI REALTIME (CAMERA) =====

        # Giữ chuột -> khung XANH DA TRỜI
        if ROI_DRAWING and not ROI_ACTIVE and roi_start and roi_end:
            cv2.rectangle(frame, roi_start, roi_end, (255,255,0), 2)

        # Đã thả chuột → khung VÀNG + chữ
        if ROI_BOX and ROI_DRAWING:
            x1, y1, x2, y2 = ROI_BOX
            cv2.rectangle(
                frame,
                (x1, y1),
                (x2, y2),
                (0, 255, 255) if not ROI_ACTIVE else (0, 255, 0),  # Vàng
                2
            )

            cv2.putText(
                frame,
                f"ROI ACTIVE | HS-{ZONE_ID}" if ROI_ACTIVE else "ROI DRAWN",
                (x1, y1 - 10),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.7,
                (0, 255, 255) if not ROI_ACTIVE else (0, 255, 0),
                2
            )

        cv2.imshow(window_title, cv2.resize(frame, (new_w, new_h)))


        camera_icon = os.path.join(BASE_DIR, "Emotion + Posture Detector v3.0 Camera.ico")

        if first_show:
            bring_window_to_front(window_title)
            set_opencv_window_icon(window_title, camera_icon)
            first_show = False

        cv2.namedWindow(window_title)
        cv2.setMouseCallback(window_title, mouse_draw_roi)

        if first_show:
            bring_window_to_front(window_title)
            first_show = False
        else:
            # Vẫn giữ cửa sổ CV2 luôn trên cùng (nếu có)
            hwnd = win32gui.FindWindow(None, window_title)
            if hwnd: win32gui.SetWindowPos(hwnd, win32con.HWND_TOPMOST, 0, 0, 0, 0, win32con.SWP_NOMOVE | win32con.SWP_NOSIZE)
        
    # KHI VÒNG LẶP KẾT THÚC: Dọn dẹp
    cap.release()
    cv2.destroyAllWindows()
    with thread_lock:
        is_running = False
        detection_thread = None # Đặt lại luồng để có thể chạy lại

    if len(DATA_LOGS) > 1 and not force_exit_no_report:
        root.after(100, analyze_and_export_csv) # Chạy hàm xuất CSV trên luồng chính Tkinter

# HÀM CHÍNH CHO FULLSCREEN

def run_detection_fullscreen():
    global latest_frame, frame_lock, is_running, root, broadcast_thread, detection_thread
    global cap, DATA_LOGS, SCAN_MIN_DURATION, WIDTH_SCR, HEIGHT_SCR # THÊM WIDTH_SCR, HEIGHT_SCR
    global INCIDENT_STATE, INCIDENT_START_TIME, INCIDENT_START_TIME_STR
    global ROI_ACTIVE, ROI_BOX, ROI_DRAWING, ROI_IMAGE_PATH
    global roi_status_color, roi_status_text, roi_emotion_label, ABNORMAL_THRESHOLD
    global roi_start, roi_end, scale_factor
    global class_name, ZONE_ID
    global force_exit_no_report

    ROI_ACTIVE = False
    ROI_DRAWING = False
    roi_start = None
    roi_end = None

    INCIDENT_STATE = None
    INCIDENT_START_TIME = None
    INCIDENT_START_TIME_STR = None

    DATA_LOGS = [] # Xóa log cũ
    scan_start_time = time.time()
    last_log_time = time.time()

    window_title = 'Emotion + Posture Detector v5.0 (Fullscreen Capture)'
    
    # KHI HÀM BẮT ĐẦU CHẠY: Báo cho GUI biết là detection đang chạy
    with thread_lock:
        is_running = True
    
    local_ip = get_local_ip()
    link = f"http://{local_ip}:5000/"
    
    if not hasattr(run_detection_fullscreen, "_flask_started"):
        Thread(target=start_flask_server, daemon=True).start()
        time.sleep(1)
        run_detection_fullscreen._flask_started = True

    if broadcast_thread is None or not broadcast_thread.is_alive():
        broadcast_thread = Thread(target=udp_broadcast, args=(link,), daemon=True)
        broadcast_thread.start()

    # show_loading_window("Đang khởi động Fullscreen Capture...")
    update_progress(25, "Đang tải mô hình nhận diện biểu cảm (Keras)...")
    import mediapipe as mp
    try:
        from tensorflow.keras.models import load_model
    except ImportError:
        messagebox.showerror("Lỗi", "Vui lòng cài đặt Tensorflow/Keras.")
        root.after(0, destroy_loading_window)
        with thread_lock:
             is_running = False
        return

    face_xml = os.path.join(BASE_DIR, "haarcascade_frontalface_default.xml")
    model_h5 = os.path.join(BASE_DIR, "emotion_detection.h5")
    face_classifier = cv2.CascadeClassifier(face_xml)
    if face_classifier.empty():
        messagebox.showerror("Lỗi", f"Không tìm thấy file cascade: {face_xml}")
        root.after(0, destroy_loading_window)
        with thread_lock:
             is_running = False
        return
    classifier = load_model(model_h5)
    
    update_progress(50, "Đang tải mô hình tư thế (MediaPipe)...")
    class_labels = ['Giận dữ', 'Ghê sợ', 'Sợ hãi', 'Vui vẻ', 'Buồn', 'Bất ngờ', 'Trung lập']
    mp_pose = mp.solutions.pose
    pose = mp_pose.Pose(min_detection_confidence=0.5, min_tracking_confidence=0.5)
    mp_drawing = mp.solutions.drawing_utils

    update_progress(100, "Hoàn tất! Mở màn hình...")
    root.after(0, destroy_loading_window)
    root.after(100, show_stream_link, link)

    WIDTH_SCR, HEIGHT_SCR = pyautogui.size() # Lấy kích thước màn hình
    scale_factor = 0.5
    first_show = True

    global session_start_time, bad_posture_total_frames, total_detection_frames, history
    session_start_time = time.time()
    bad_posture_total_frames = 0
    total_detection_frames = 0
    history.clear() # Đảm bảo history sạch khi bắt đầu phiên mới

    start_time = time.time()
    interval = 120

    font = ImageFont.truetype(font_path, 28)
    font2 = ImageFont.truetype(font_path, 20)

    short_term_emotion_buffer = deque(maxlen=STABILITY_WINDOW_FRAMES)
    bad_posture_counter = 0 
    current_stable_emotion = 'Trung lập'

    status_posture = "Không phát hiện tư thế"
        
    force_exit_no_report = False
    # Vòng lặp chính
    while is_running:

        current_time = time.time()
        frame = np.array(pyautogui.screenshot())
        frame = cv2.cvtColor(frame, cv2.COLOR_RGB2BGR)

        # === BẮT ĐẦU: Logic che cửa sổ OpenCV ===
        hwnd = win32gui.FindWindow(None, window_title)
        if hwnd:
            try:
                rect = win32gui.GetWindowRect(hwnd)
                x_win, y_win, x_end_win, y_end_win = rect
                # Đảm bảo tọa độ hợp lệ
                x_win, y_win = max(0, x_win), max(0, y_win)
                x_end_win, y_end_win = min(WIDTH_SCR, x_end_win), min(HEIGHT_SCR, y_end_win)
                
                # 2. Vẽ đè (làm đen) khu vực cửa sổ OpenCV
                if x_end_win > x_win and y_end_win > y_win:
                    # Chụp màn hình thường có kênh màu BGR, không phải RGB
                    frame[y_win:y_end_win, x_win:x_end_win] = (0, 0, 0) # Màu đen BGR
            except Exception:
                # Bỏ qua nếu có lỗi khi lấy tọa độ cửa sổ
                pass
        # === KẾT THÚC: Logic che cửa sổ OpenCV ===
        
        # --- Emotion Detection ---
        frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(frame_pil)
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        faces = face_classifier.detectMultiScale(gray, 1.3, 5)
        
        final_emotion_label = current_stable_emotion
        roi_emotion_label = current_stable_emotion

        for (x, y, w, h) in faces:
            face_cx = x + w // 2
            face_cy = y + h // 2

            face_in_roi = False
            if ROI_ACTIVE and ROI_BOX:
                rx1, ry1, rx2, ry2 = ROI_BOX
                if rx1 <= face_cx <= rx2 and ry1 <= face_cy <= ry2:
                    face_in_roi = True
            cv2.rectangle(frame, (x, y), (x + w, y + h), (242, 248, 68), 2)
            roi_gray = cv2.resize(gray[y:y + h, x:x + w], (48, 48), interpolation=cv2.INTER_AREA)
            roi = np.expand_dims(np.expand_dims(roi_gray.astype("float") / 255.0, axis=-1), axis=0)
            
            # Lấy xác suất
            predictions = classifier.predict(roi, verbose=0)
            probabilities = predictions[0] 
            sorted_indices = np.argsort(probabilities)[::-1]
            
            p_max = probabilities[sorted_indices[0]]
            p_top2 = probabilities[sorted_indices[1]]
            predicted_label = class_labels[sorted_indices[0]] # Nhãn dự đoán

            # 1. KIỂM TRA NGƯỠNG (VALIDATION)
            config = EMOTION_THRESHOLDS.get(predicted_label, {'p_max': THRESHOLD_P_MAX_DEFAULT, 'delta': THRESHOLD_DELTA_TOP2_DEFAULT})
            
            is_reliable = p_max >= config['p_max']
            is_not_ambiguous = (p_max - p_top2) >= config['delta']
            
            if is_reliable and is_not_ambiguous:
                raw_validated_label = predicted_label
            else:
                raw_validated_label = 'Unknown' # Nhãn không đủ tin cậy

            # 2. CỬA SỔ ỔN ĐỊNH NGẮN HẠN (W: 2-3 GIÂY)
            short_term_emotion_buffer.append(raw_validated_label)
            
            if len(short_term_emotion_buffer) == STABILITY_WINDOW_FRAMES:
                counts = Counter(short_term_emotion_buffer)
                most_common_label, count = counts.most_common(1)[0]
                
                # Nếu nhãn phổ biến nhất chiếm ưu thế (>= 80%) và không phải Unknown
                if most_common_label != 'Unknown' and (count / STABILITY_WINDOW_FRAMES) >= STABILITY_DOMINANCE_RATIO: 
                    final_emotion_label = most_common_label
                # Ngược lại, giữ nhãn ổn định trước đó
                else: 
                    final_emotion_label = current_stable_emotion 
                
                # Cập nhật nhãn ổn định hiện tại và lưu vào lịch sử dài hạn
                current_stable_emotion = final_emotion_label
                if current_stable_emotion != 'Unknown':
                     # 1 = Tiêu cực; 0 = Tích cực/Trung lập
                     if current_stable_emotion in ['Giận dữ', 'Ghê sợ', 'Sợ hãi', 'Buồn']:
                         history.append(1)
                     else:
                         history.append(0)
                
                # Xóa buffer để bắt đầu cửa sổ mới
                short_term_emotion_buffer.clear()
            else:
                final_emotion_label = current_stable_emotion # Hiển thị nhãn ổn định trước đó
            
            if face_in_roi:
                roi_emotion_label = final_emotion_label
            # Vẽ nhãn cảm xúc đã được ổn định lên frame
            frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
            draw = ImageDraw.Draw(frame_pil)
            draw_text_with_outline(draw, (x, y - 35), final_emotion_label, font, (0, 255, 0))
            frame = cv2.cvtColor(np.array(frame_pil), cv2.COLOR_RGB2BGR)
        
        negative_ratio = sum(history) / len(history) if len(history) > 0 else 0
        elapsed = time.time() - start_time
        if elapsed >= interval:
            if negative_ratio > 0.6:
                show_warning("Pause / đổi hoạt động / nghỉ 2 phút")
            start_time = time.time()

        labels = [final_emotion_label]
        # --- Posture Detection (Đã đồng bộ) ---
        if ROI_ACTIVE and ROI_BOX:
            x1, y1, x2, y2 = ROI_BOX

            # Cắt ảnh theo ROI
            roi_frame = frame[y1:y2, x1:x2]

            if roi_frame.size == 0:
                results = None
            else:
                image_rgb = cv2.cvtColor(roi_frame, cv2.COLOR_BGR2RGB)
                results = pose.process(image_rgb)
        else:
            image_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
            results = pose.process(image_rgb)
        angle_back = 0
        status_back_detail = "Không phát hiện tư thế"
        status_posture = "Không phát hiện tư thế"
        color = (255, 255, 255)
        
        if results is not None and results.pose_landmarks:
            landmarks = results.pose_landmarks.landmark
            
            person_in_roi = False

            if ROI_ACTIVE and ROI_BOX:
                x1, y1, x2, y2 = ROI_BOX
                roi_w = x2 - x1
                roi_h = y2 - y1

                # Landmark đang là tọa độ TRONG ROI
                cx = int(
                    ((landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].x +
                      landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].x) * 0.5) * roi_w
                ) + x1

                cy = int(
                    ((landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].y +
                      landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].y) * 0.5) * roi_h
                ) + y1

                if x1 <= cx <= x2 and y1 <= cy <= y2:
                    person_in_roi = True

            if (not ROI_ACTIVE) or person_in_roi:
                try:
                    ear_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].y
                    ]
                    shoulder_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_SHOULDER.value].y
                    ]
                    hip_l = [
                        landmarks[mp_pose.PoseLandmark.LEFT_HIP.value].x,
                        landmarks[mp_pose.PoseLandmark.LEFT_HIP.value].y
                    ]

                    ear_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].y
                    ]
                    shoulder_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_SHOULDER.value].y
                    ]
                    hip_r = [
                        landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value].x,
                        landmarks[mp_pose.PoseLandmark.RIGHT_HIP.value].y
                    ]

                    vis_l = landmarks[mp_pose.PoseLandmark.LEFT_EAR.value].visibility
                    vis_r = landmarks[mp_pose.PoseLandmark.RIGHT_EAR.value].visibility

                    angle_back = None

                    if vis_l >= vis_r and vis_l > 0.5:
                        angle_back = calculate_angle(hip_l, shoulder_l, ear_l)
                    elif vis_r > 0.5:
                        angle_back = calculate_angle(hip_r, shoulder_r, ear_r)
                    else:
                        angle_back = None  # Không đủ tin cậy


                    # --- Quy đổi Ergonomics (Dựa trên Góc Lưng) ---
                    if angle_back is None:
                        status_posture = "Không đủ dữ liệu"
                        color = (255, 255, 255)

                    else:
                        if angle_back >= 170:
                            status_back_detail = "Ngồi thẳng (Good)"
                            status_posture = "Ngồi thẳng (Good)"
                            color = (0, 255, 0) # Xanh
                        elif 150 <= angle_back < 170:
                            status_back_detail = "Hơi cúi (Warning)"
                            status_posture = "Hơi cúi (Warning)"
                            color = (255, 255, 0) # Vàng
                        else:
                            status_back_detail = "Cúi nhiều (Bad)"
                            status_posture = "Cúi nhiều (Bad)"
                            color = (255, 0, 0) # Đỏ
                    
                    # Vẽ Landmarks
                    if ROI_ACTIVE and ROI_BOX:
                        mp_drawing.draw_landmarks(
                            roi_frame,
                            results.pose_landmarks,
                            mp_pose.POSE_CONNECTIONS
                        )
                        # dán ROI đã vẽ landmark lại frame gốc
                        frame[y1:y2, x1:x2] = roi_frame
                    else:
                        mp_drawing.draw_landmarks(
                            frame,
                            results.pose_landmarks,
                            mp_pose.POSE_CONNECTIONS
                        )
                except Exception:
                    # Xử lý khi không tìm thấy đủ 3 điểm (Xảy ra khi quay nghiêng quá nhiều)
                    status_back_detail = "Không đủ điểm (Quá nghiêng)"
                    status_posture = "Không phát hiện tư thế"
                    color = (255, 255, 255) # Trắng

            if ROI_ACTIVE and person_in_roi:
                # Cho file word
                ROI_LOGS.append({ 
                        "time": time.time(), 
                        "emotion": roi_emotion_label, 
                        "posture": status_posture 
                    })

                # Cho trang web
                current_state = f"{roi_emotion_label}/{status_posture}"
                now = time.time()

                # Nếu chưa có trạng thái → khởi tạo
                if ROI_STATE_TRACKER["state"] is None:
                    ROI_STATE_TRACKER["state"] = current_state
                    ROI_STATE_TRACKER["start_time"] = now

                # Nếu trạng thái THAY ĐỔI
                elif current_state != ROI_STATE_TRACKER["state"]:
                    prev_state = ROI_STATE_TRACKER["state"]
                    start_time_roi = ROI_STATE_TRACKER["start_time"]
                    duration = int(now - start_time_roi)

                    # Kiểm tra trạng thái cũ có bất thường không
                    is_abnormal = (
                        "Warning" in prev_state or "Bad" in prev_state or
                        any(x in prev_state for x in ["Buồn", "Giận dữ", "Ghê sợ", "Sợ hãi"])
                    )

                    # Nếu bất thường và đủ thời gian → GỬI NGAY
                    if is_abnormal and duration >= ABNORMAL_THRESHOLD:
                        now_dt = datetime.datetime.now()
                        start_time_str = (now_dt - datetime.timedelta(seconds=duration)).strftime('%H:%M:%S')

                        send_incident(
                            prev_state,
                            start_time_str,
                            duration
                        )

                    # Reset sang trạng thái mới
                    ROI_STATE_TRACKER["state"] = current_state
                    ROI_STATE_TRACKER["start_time"] = now

        now_time = time.time()
        now = datetime.datetime.now()

        combined_state = f"{final_emotion_label}/{status_posture}"

        is_abnormal = (
            final_emotion_label in ['Buồn', 'Giận dữ', 'Sợ hãi', 'Ghê sợ']
            or "Bad" in status_posture
            or "Warning" in status_posture
        )

        if ROI_ACTIVE and is_abnormal:

            # === TRẠNG THÁI MỚI ===
            if INCIDENT_STATE != combined_state:

                # Nếu có trạng thái cũ → kiểm tra gửi
                if INCIDENT_STATE is not None:
                    duration = int(now_time - INCIDENT_START_TIME)
                    if duration >= ABNORMAL_THRESHOLD:
                        send_incident(
                            INCIDENT_STATE,
                            INCIDENT_START_TIME_STR,
                            duration
                        )

                # Bắt đầu trạng thái mới
                INCIDENT_STATE = combined_state
                INCIDENT_START_TIME = now_time
                INCIDENT_START_TIME_STR = now.strftime('%H:%M:%S')

            else:
                # === TRẠNG THÁI GIỮ NGUYÊN ===
                duration = int(now_time - INCIDENT_START_TIME)
                if duration >= ABNORMAL_THRESHOLD:
                    send_incident(
                        INCIDENT_STATE,
                        INCIDENT_START_TIME_STR,
                        duration
                    )

                    # reset để tránh spam
                    INCIDENT_START_TIME = now_time
                    INCIDENT_START_TIME_STR = now.strftime('%H:%M:%S')

        else:
            # === HẾT BẤT THƯỜNG → gửi nốt nếu cần ===
            if INCIDENT_STATE is not None:
                duration = int(now_time - INCIDENT_START_TIME)
                if duration >= ABNORMAL_THRESHOLD:
                    send_incident(
                        INCIDENT_STATE,
                        INCIDENT_START_TIME_STR,
                        duration
                    )

            INCIDENT_STATE = None
            INCIDENT_START_TIME = None
            INCIDENT_START_TIME_STR = None

        # --- LOGIC CẢNH BÁO TƯ THẾ ĐỘC LẬP (Đã đồng bộ) ---
        if "Bad" in status_posture:
            bad_posture_counter += 1
        else:
            bad_posture_counter = 0

        if bad_posture_counter >= BAD_POSTURE_WARNING_FRAMES:
            show_warning("CẢNH BÁO TƯ THẾ: Bạn đã ngồi cúi gù quá lâu. Vui lòng điều chỉnh lại tư thế ngồi!")
            bad_posture_counter = 0 # Reset sau cảnh báo

        total_detection_frames += 1 # Tăng tổng frame

        box_color = (0, 0, 255) if negative_ratio > 0.6 else (0, 255, 255) if 0.2 <= negative_ratio <= 0.6 else (0, 255, 0)


        # --- LOGGING DỮ LIỆU ---
        # Ghi log mỗi 0.5 giây để tiết kiệm tài nguyên
        if current_time - last_log_time >= 0.5 and labels:
            with LOG_LOCK:
                # SỬ DỤNG status_posture ĐÃ ĐƯỢC XÁC ĐỊNH
                DATA_LOGS.append({
                    'timestamp': current_time,
                    'emotion': labels[0] if labels else 'Không phát hiện', 
                    'posture_status': status_posture,
                })
            last_log_time = current_time

        # --- DRAWING TEXT ---
        frame_pil = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(frame_pil)

        # --- Hiển thị Text ---
        draw_text_with_outline(draw, (20, 20), f"Tư thế Lưng: {status_back_detail} ({int(angle_back)}°)", font2, (255, 255, 255))
        draw_text_with_outline(draw, (20, 50), f"Trạng thái tổng: {status_posture}", font2, color) # Hiển thị tóm tắt và màu
        draw_text_with_outline(draw, (20, 80), f"Số lượng: {len(faces)}", font2, (255, 0, 255))
        draw_text_with_outline(draw, (20, 110), "Trạng thái:", font2, (0, 0, 255))
        draw_text_with_outline(draw, (int(WIDTH_SCR * 0.75 + 145), HEIGHT_SCR - 40), "Bấm phím 'Q' để thoát", font, (255, 255, 0)) # Chú thích thay đổi
        draw_text_with_outline(draw, (int(WIDTH_SCR * 0.8 + 90), 20), "Bấm phím 'M' để phóng to", font2, (255, 255, 0))
        draw_text_with_outline(draw, (int(WIDTH_SCR * 0.8 + 90), 50), "Bấm phím 'N' để thu nhỏ", font2, (255, 255, 0))
        draw_text_with_outline(draw, (int(WIDTH_SCR * 0.8 + 90), 80), "Đang quét tại: " + class_name, font2, (0, 255, 0))
        draw_text_with_outline(draw, (int(WIDTH_SCR * 0.8 + 90), 110), roi_status_text, font2, roi_status_color)
        frame = cv2.cvtColor(np.array(frame_pil), cv2.COLOR_RGB2BGR)

        # Draw box màu (Điều chỉnh tọa độ cho phù hợp)
        draw_filled_rectangle_with_outline(frame, (235 - 100, 125 - 13), (255 - 100, 145 - 13), box_color, outline_width=2) 
        
        # ... (Phần xử lý phím bấm và hiển thị giữ nguyên) ...
        key = cv2.waitKey(10) & 0xFF
        if key == ord('m'): scale_factor = min(0.9, scale_factor + 0.1)
        elif key == ord('n'): scale_factor = max(0.2, scale_factor - 0.1)
        elif key == ord('q'): 
            # Bắt buộc phải quét tối thiểu 30s
            current_time = time.time()
            elapsed = current_time - scan_start_time

            if elapsed < SCAN_MIN_DURATION:
                remaining = int(SCAN_MIN_DURATION - elapsed)

                answer = ask_yes_no_on_main_thread(
                    "Chưa đủ thời gian quét",
                    f"Cần quét tối thiểu {SCAN_MIN_DURATION} giây.\n"
                    f"Bạn cần quét thêm {remaining} giây nữa.\n\n"
                    f"Nếu dừng bây giờ sẽ KHÔNG xuất báo cáo.\n"
                    f"Bạn có chắc chắn muốn dừng không?"
                )

                if answer:
                    force_exit_no_report = True
                    break
                else:
                    continue
            else:
                break


        elif key == ord('v') and not ROI_ACTIVE:
            ROI_DRAWING = not ROI_DRAWING

            # CHỈ xoá khung khi TẮT vẽ
            if not ROI_DRAWING:
                roi_start = None
                roi_end = None
                ROI_BOX = None

        elif key == ord('s') and ROI_BOX and not ROI_ACTIVE:
            global ROI_IMAGE_BUFFER
            if not ZONE_ID:
                ok = ask_student_id(root)

                # ÉP cửa sổ OpenCV hiện lại sau khi đóng Tkinter
                cv2.namedWindow(window_title)
                hwnd = win32gui.FindWindow(None, window_title)
                if hwnd:
                    win32gui.SetWindowPos(
                        hwnd,
                        win32con.HWND_TOPMOST,
                        0, 0, 0, 0,
                        win32con.SWP_NOMOVE | win32con.SWP_NOSIZE
                    )

                if not ok:
                    show_warning("Vui lòng nhập Student ID trước khi quét ROI.")
                    continue

            ROI_ACTIVE = True
            ROI_LOGS.clear()
            roi_scan_start_time = time.time()

            x1, y1, x2, y2 = ROI_BOX
            x1, y1 = max(0, x1), max(0, y1)
            x2, y2 = min(frame.shape[1], x2), min(frame.shape[0], y2)

            roi_crop = frame[y1:y2, x1:x2]

            if roi_crop.size > 0:
                _, buffer = cv2.imencode(".png", roi_crop)
                ROI_IMAGE_BUFFER = buffer.tobytes()

            show_warning("BẮT ĐẦU QUÉT ROI (ĐÃ CHỤP ẢNH ROI)")


        elif key == ord('e') and ROI_ACTIVE:
            # GỬI TRẠNG THÁI CUỐI CÙNG NẾU ĐỦ ĐIỀU KIỆN
            if ROI_STATE_TRACKER["state"] and ROI_STATE_TRACKER["start_time"]:
                now = time.time()
                duration = now - ROI_STATE_TRACKER["start_time"]

                if duration >= ABNORMAL_THRESHOLD:
                    duration = int(time.time() - ROI_STATE_TRACKER["start_time"])

                    now_dt = datetime.datetime.now()
                    start_time_str = (now_dt - datetime.timedelta(seconds=duration)).strftime('%H:%M:%S')

                    send_incident(
                        ROI_STATE_TRACKER["state"],
                        start_time_str,
                        duration
                    )

            ROI_ACTIVE = False
            ROI_STATE_TRACKER["state"] = None
            ROI_STATE_TRACKER["start_time"] = None

            export_roi_to_word()
            ROI_IMAGE_PATH = None
            ZONE_ID = None

        if ROI_DRAWING:
            roi_status_text = "Vẽ khung ROI: BẬT"
            roi_status_color = (0, 255, 0)
        else:
            roi_status_text = "Vẽ khung ROI: TẮT"
            roi_status_color = (255, 0, 0)
        frame_stream = cv2.resize(frame.copy(), (int(WIDTH_SCR*0.4), int(HEIGHT_SCR*0.4)))
        with frame_lock:
            latest_frame = frame_stream.copy()
            
        new_w, new_h = int(WIDTH_SCR * scale_factor), int(HEIGHT_SCR * scale_factor)

        # ===== VẼ ROI REALTIME (FULLSCREEN) =====

        # Đang kéo chuột → khung XANH DA TRỜI
        if ROI_DRAWING and not ROI_ACTIVE and roi_start and roi_end:
            cv2.rectangle(frame, roi_start, roi_end, (255,255,0), 2)

        # Đã thả chuột → khung VÀNG + chữ
        if ROI_BOX and ROI_DRAWING:
            x1, y1, x2, y2 = ROI_BOX
            cv2.rectangle(
                frame,
                (x1, y1),
                (x2, y2),
                (0, 255, 255) if not ROI_ACTIVE else (0, 255, 0),  # Vàng
                2
            )

            cv2.putText(
                frame,
                f"ROI ACTIVE | HS-{ZONE_ID}" if ROI_ACTIVE else "ROI DRAWN",
                (x1, y1 - 10),
                cv2.FONT_HERSHEY_SIMPLEX,
                0.7,
                (0, 255, 255) if not ROI_ACTIVE else (0, 255, 0),
                2
            )

        cv2.imshow(window_title, cv2.resize(frame, (new_w, new_h)))
        
        fullscreen_icon = os.path.join(
            BASE_DIR,
            "Emotion + Posture Detector v3.0 Fullscreen Capture.ico"
        )

        if first_show:
            bring_window_to_front(window_title)
            set_opencv_window_icon(window_title, fullscreen_icon)
            first_show = False

        cv2.namedWindow(window_title)
        cv2.setMouseCallback(window_title, mouse_draw_roi_fullscreen)

        if first_show:
            bring_window_to_front(window_title)
            first_show = False
        else:
            # Vẫn giữ cửa sổ CV2 luôn trên cùng (nếu có)
            hwnd = win32gui.FindWindow(None, window_title)
            if hwnd: win32gui.SetWindowPos(hwnd, win32con.HWND_TOPMOST, 0, 0, 0, 0, win32con.SWP_NOMOVE | win32con.SWP_NOSIZE)
        
    # KHI VÒNG LẶP KẾT THÚC: Dọn dẹp
    cv2.destroyAllWindows()
    with thread_lock:
        is_running = False
        detection_thread = None # Đặt lại luồng để có thể chạy lại

    if len(DATA_LOGS) > 1 and not force_exit_no_report:
        root.after(100, analyze_and_export_csv) # Chạy hàm xuất CSV trên luồng chính Tkinter


# --- GIAO DIỆN CHÍNH (GUI CHỌN LỰA) ---

def select_camera_and_run(cam_index):
    global root, current_mode, detection_thread
    if cam_index == -1:
        messagebox.showwarning("Chưa chọn", "Vui lòng chọn một camera.")
        return
    
    with thread_lock:
        if is_running:
            messagebox.showwarning("Đang chạy", "Một chế độ quét đang chạy. Vui lòng tắt cửa sổ quét (hoặc nhấn 'Q') trước.")
            return

    current_mode = 'camera'
    # KHÔNG DÙNG root.withdraw() nữa

    # Hiển thị loading trước khi chạy luồng detection
    show_loading_window("Đang khởi động Camera...")
    
    # Chạy detection trong luồng riêng
    detection_thread = Thread(target=run_detection_camera, args=(cam_index,), daemon=True)
    detection_thread.start()

def open_camera_selection_dialog():
    """Hiển thị cửa sổ chọn camera và lớp trước khi chạy chế độ Camera."""
    global cam_window, camera_combo, class_combo, class_name, root, icon_path_camera

    # Kiểm tra xem có đang chạy chế độ nào không
    with thread_lock:
        if is_running:
            messagebox.showwarning("Đang chạy", 
                                   "Một chế độ quét đang chạy. Vui lòng tắt cửa sổ quét (hoặc nhấn 'Q') trước.")
            return 
    
    cam_window = tk.Toplevel(root)
    cam_window.title("Chọn Camera và Lớp")
    if os.path.exists(icon_path_camera):
        cam_window.iconbitmap(icon_path_camera)
    # cam_window.attributes('-topmost', True)
    
    cameras = ["Chọn Camera..."] + list_cameras()
    tk.Label(cam_window, text="Vui lòng chọn camera để quét:", font=("Arial", 10)).pack(pady=5)

    camera_combo = ttk.Combobox(cam_window, values=cameras, state="readonly", width=30)
    camera_combo.current(0)
    camera_combo.pack(pady=5)

    # --- Thêm ComboBox chọn lớp ---
    tk.Label(cam_window, text="Chọn lớp học:", font=("Arial", 10)).pack(pady=5)
    classes = ["Lớp 12A1", "Lớp 12A2", "Lớp 12A3", "Lớp 12A4", "Lớp 12A5"]
    class_combo = ttk.Combobox(cam_window, values=classes, state="readonly", width=30)
    class_combo.current(0)  # Mặc định chọn lớp đầu tiên
    class_combo.pack(pady=5)

    tk.Label(cam_window, text="Hướng dẫn vẽ khung:", font=("Arial", 8)).pack(pady=1)
    tk.Label(cam_window, text="1. Bấm phím V để bật/tắt vẽ khung", font=("Arial", 8)).pack(pady=1)
    tk.Label(cam_window, text="2. Bấm phím S để bắt đầu quét", font=("Arial", 8)).pack(pady=1)
    tk.Label(cam_window, text="3. Bấm phím E để dừng quét và xuất file", font=("Arial", 8)).pack(pady=1)

    def on_run():
        selected_index = camera_combo.current() - 1
        if selected_index < 0:
            messagebox.showwarning("Chưa chọn", "Vui lòng chọn một camera trong danh sách.")
            return
        
        # Lấy lớp học được chọn
        global class_name
        class_name = class_combo.get()
        if not class_name:
            messagebox.showwarning("Chưa chọn", "Vui lòng chọn một lớp học.")
            return

        select_camera_and_run(selected_index)

    tk.Button(cam_window, text="Mở camera", command=on_run,
              bg="#007ACC", fg="white", font=("Arial", 10, "bold")).pack(pady=10)

    # Căn giữa cửa sổ
    root.update_idletasks()
    cam_window.update_idletasks()
    x = root.winfo_x() + (root.winfo_width() - cam_window.winfo_reqwidth()) // 2
    y = root.winfo_y() + (root.winfo_height() - cam_window.winfo_reqheight()) // 2
    cam_window.geometry(f"+{x}+{y}")


def open_class_selection_dialog_for_fullscreen(on_confirm):
    """Hiển thị hộp thoại chọn lớp trước khi chạy Fullscreen"""
    global root, class_name

    class_window = tk.Toplevel(root)
    class_window.title("Chọn lớp")
    if os.path.exists(icon_path_fullscreen):
        class_window.iconbitmap(icon_path_fullscreen)
    class_window.resizable(False, False)
    class_window.grab_set()  # Khóa các cửa sổ khác

    tk.Label(
        class_window,
        text="Vui lòng chọn lớp:",
        font=("Arial", 10, "bold")
    ).pack(pady=(12, 6))

    classes = [
        "Chọn lớp...",
        "Lớp 12A1",
        "Lớp 12A2",
        "Lớp 12A3",
        "Lớp 12A4",
        "Lớp 12A5"
    ]

    class_combo = ttk.Combobox(
        class_window,
        values=classes,
        state="readonly",
        width=30
    )
    class_combo.current(0)
    class_combo.pack(pady=5)

    def confirm():
        global class_name
        selected = class_combo.get()

        if selected == "Chọn lớp...":
            messagebox.showwarning(
                "Chưa chọn",
                "Vui lòng chọn một lớp trước khi tiếp tục."
            )
            return

        class_name = selected  # GÁN BIẾN TOÀN CỤC
        class_window.destroy()

        if callable(on_confirm):
            on_confirm()

    tk.Button(
        class_window,
        text="Xác nhận",
        command=confirm,
        bg="#007ACC",
        fg="white",
        font=("Arial", 10, "bold"),
        width=18
    ).pack(pady=12)

    # ===== Căn giữa cửa sổ =====
    root.update_idletasks()
    class_window.update_idletasks()

    x = root.winfo_x() + (root.winfo_width() - class_window.winfo_reqwidth()) // 2
    y = root.winfo_y() + (root.winfo_height() - class_window.winfo_reqheight()) // 2

    class_window.geometry(f"+{x}+{y}")

def start_fullscreen_capture():
    global root, current_mode, detection_thread
    
    with thread_lock:
        if is_running:
             messagebox.showwarning("Đang chạy", "Một chế độ quét đang chạy. Vui lòng tắt cửa sổ quét (hoặc nhấn 'Q') trước.")
             return 
    
    def after_class_selected():
        global current_mode, detection_thread

        current_mode = 'screen'

        show_loading_window("Đang khởi động Fullscreen Capture...")

        # Chạy detection trong luồng riêng
        detection_thread = Thread(
            target=run_detection_fullscreen,
            daemon=True
        )
        detection_thread.start()

    open_class_selection_dialog_for_fullscreen(after_class_selected)



def stop_detection():
    # Chỉ giữ lại các biến global cần thiết cho việc dừng luồng và reset
    global is_running, detection_thread, broadcast_thread, root

    # 1. Dừng luồng quét
    with thread_lock:
        is_running = False
        
    # 2. GỌI analyze_and_export_csv trên luồng chính Tkinter
    # Dù stop_detection được gọi từ đâu, ta dùng root.after để đảm bảo an toàn cho Tkinter calls
    if root:
        root.after(100, analyze_and_export_csv) 
        
    # 3. Dọn dẹp luồng 
    detection_thread = None
    broadcast_thread = None

    print("Quá trình dừng đã hoàn tất.")

# Sửa luôn hàm on_closing() để nó gọi stop_detection()
def on_closing():
    global is_running, root
    with thread_lock:
        is_running_local = is_running

    if is_running_local:
        if messagebox.askyesno("Thoát", "Chương trình quét đang chạy. Bạn có muốn dừng và thoát không?"):
            stop_detection() # Dừng luồng quét (sẽ tự động gọi analyze_and_export_csv)
            root.quit() # Thoát khỏi Tkinter
    else:
        root.quit()

root = tk.Tk()
root.title("Lựa Chọn Chế Độ")
if os.path.exists(icon_path):
    root.iconbitmap(icon_path)
    
# Gán hàm xử lý sự kiện đóng cửa sổ (Sửa lỗi ở đây)
root.protocol("WM_DELETE_WINDOW", on_closing) 

window_width = 550
# Chiều cao đã được tăng lên để chứa 4 nút
window_height = 300
screen_width = root.winfo_screenwidth()
screen_height = root.winfo_screenheight()
x = (screen_width // 2) - (window_width // 2)
y = (screen_height // 2) - (window_height // 2)
root.geometry(f'{window_width}x{window_height}+{x}+{y}')
root.resizable(False, False)


tk.Label(root, text="CHỌN CHẾ ĐỘ QUÉT BIỂU CẢM KHUÔN MẶT/TƯ THẾ:", 
         font=("Arial", 12, "bold")).pack(pady=10)

tk.Button(root, text="📷 QUÉT BẰNG CAMERA", 
          command=open_camera_selection_dialog,
          width=30, height=2, bg="#007ACC", fg="white", font=("Arial", 10, "bold")).pack(pady=5)

tk.Button(root, text="🖥️ QUÉT BẰNG MÀN HÌNH",
          command=start_fullscreen_capture,
          width=30, height=2, bg="#FF9800", fg="white", font=("Arial", 10, "bold")).pack(pady=5)

tk.Button(root, text="AI Smart Monitor",
          command=open_aismartmonitor,
          width=30, height=2, bg="#216C71", fg="white", font=("Arial", 10, "bold")).pack(pady=5)

# Nút tùy chọn thư mục log
tk.Button(root, text="📁 Xuất file log tại...",
          command=set_log_directory,
          width=30, height=2, bg="#38505D", fg="white", font=("Arial", 10, "bold")).pack(pady=5)

root.mainloop()
