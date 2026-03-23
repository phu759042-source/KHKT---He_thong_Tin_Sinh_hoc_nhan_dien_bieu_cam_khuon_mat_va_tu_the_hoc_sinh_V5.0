import sys
import os
import cv2
import json
import time
import datetime
import numpy as np
import tkinter as tk
from tkinter import messagebox, ttk, filedialog
import traceback
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if sys.platform == "win32":
    try:
        import ctypes
        ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(
            "AISmartMonitor.AISmartMonitor.App"
        )
    except Exception as e:
        print("Không set được AppUserModelID:", e)


SERVER_URL = "https://epd-test.onrender.com/log_incident/"

# --- LOADING WINDOW ---
loading_window = None
progress_bar = None
progress_label = None
root = tk.Tk()
root.withdraw()  
if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
icon_path = os.path.join(BASE_DIR, "AI Smart Monitor.ico")

def show_loading_window(title="Đang khởi động hệ thống..."):
    global loading_window, progress_bar, progress_label, root, icon_path

    loading_window = tk.Toplevel(root)
    loading_window.title(title)

    width = 400
    height = 140

    # Lấy độ phân giải màn hình
    screen_width = loading_window.winfo_screenwidth()
    screen_height = loading_window.winfo_screenheight()

    # Tính vị trí giữa màn hình
    x = (screen_width - width) // 2
    y = (screen_height - height) // 2

    # Set geometry
    loading_window.geometry(f"{width}x{height}+{x}+{y}")
    loading_window.resizable(False, False)

    # Icon
    if os.path.exists(icon_path):
        loading_window.iconbitmap(icon_path)

    # Nội dung
    tk.Label(
        loading_window,
        text="Đang khởi động hệ thống, vui lòng chờ...",
        font=("Arial", 10)
    ).pack(pady=10)

    progress_bar = ttk.Progressbar(
        loading_window,
        orient="horizontal",
        length=350,
        mode="determinate"
    )
    progress_bar.pack(pady=10)
    progress_bar["maximum"] = 100
    progress_bar["value"] = 0

    progress_label = tk.Label(
        loading_window,
        text="0%",
        font=("Arial", 10, "bold")
    )
    progress_label.pack()

    # Luôn nổi trên
    loading_window.attributes("-topmost", True)

def update_progress(percent, text=None):
    if progress_bar and progress_label and loading_window and loading_window.winfo_exists():
        progress_bar["value"] = percent
        if text: progress_label.config(text=f"{text} ({percent}%)")
        else: progress_label.config(text=f"{percent}%")
        loading_window.update_idletasks()

def destroy_loading_window():
    global loading_window
    if loading_window and loading_window.winfo_exists():
        loading_window.destroy()

# --- IMPORTS SAU LOADING ---
show_loading_window("Đang khởi động chương trình...")
update_progress(50, "Đang khởi động mô hình mediapipe")
import mediapipe as mp
update_progress(70, "Đang khởi động giao diện")
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                             QHBoxLayout, QPushButton, QLabel, QComboBox, 
                             QFrame, QSplitter, QMessageBox, QInputDialog, 
                             QFileDialog, QScrollArea, QGridLayout, QMenu)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QImage, QPixmap, QFont, QIcon
from PyQt6.QtMultimedia import QMediaDevices

update_progress(100, "Hoàn tất")
destroy_loading_window()

# --- HÀM GỬI SERVER ---
def send_to_server(data, class_name, mode):
    try:
        now = datetime.datetime.now()
        end_str = now.strftime('%H:%M:%S')
        date_str = now.strftime('%Y-%m-%d')

        payload = {
            "class_id": class_name,
            "zone_id": str(data['student_id']),
            "issue_type": data['behavior'],
            "start_time": data['timestamp'],
            "end_time": end_str,
            "duration_seconds": data['duration'],
            "date": date_str,
            "scan_mode": mode
        }
        requests.post(SERVER_URL, json=payload, timeout=2)
        print(f"✅ Đã gửi lên Web: HS-{data['student_id']} - {data['behavior']}")
    except requests.exceptions.Timeout:
        print("❌ Gửi lên web bị Timeout.")
    except Exception as e:
        print(f"❌ Lỗi khi gửi dữ liệu lên Web: {e}")

# --- BEHAVIOR TRACKER ---
class BehaviorTracker:
    def __init__(self, threshold_seconds=5):
        self.threshold = threshold_seconds
        self.tracking_data = {} 
        
    def update(self, student_id, current_status):
        now = time.time()
        if student_id not in self.tracking_data:
            if self.is_bad_behavior(current_status):
                self.tracking_data[student_id] = {'status': current_status, 'start_time': now}
            return None

        data = self.tracking_data[student_id]
        if current_status != data['status']:
            duration = now - data['start_time']
            last_status = data['status']
            start_ts = data['start_time']
            del self.tracking_data[student_id]
            
            if self.is_bad_behavior(current_status):
                self.tracking_data[student_id] = {'status': current_status, 'start_time': now}

            if duration >= self.threshold:
                return self.create_report(student_id, last_status, duration, start_ts)
        return None

    def is_bad_behavior(self, status):
        # Các từ khóa hành vi tiêu cực cần theo dõi để xếp hạng
        bad_keywords = ["Buồn ngủ", "Mất tập trung", "Căng thẳng", "Buồn", "Mệt mỏi", "Thu mình", "Kiệt sức", "Lo âu"]
        return any(keyword in status for keyword in bad_keywords)

    def create_report(self, student_id, status, duration, start_timestamp):
        return {
            "student_id": student_id,
            "behavior": status,
            "duration": round(duration, 1),
            "timestamp": datetime.datetime.fromtimestamp(start_timestamp).strftime('%H:%M:%S')
        }
    
    def finalize_all(self):
        reports = []
        now = time.time()
        for sid, data in self.tracking_data.items():
            duration = now - data['start_time']
            if duration >= self.threshold:
                reports.append(self.create_report(sid, data['status'], duration, data['start_time']))
        self.tracking_data.clear()
        return reports

# --- AI PROCESSOR ---
class AIProcessor:
    def __init__(self):
        self.mp_face_mesh = mp.solutions.face_mesh
        self.face_mesh = self.mp_face_mesh.FaceMesh(max_num_faces=1, refine_landmarks=True, min_detection_confidence=0.5, min_tracking_confidence=0.5)

    def calculate_ear(self, landmarks, indices):
        try:
            p2_p6 = np.linalg.norm(np.array([landmarks[indices[1]].x, landmarks[indices[1]].y]) - np.array([landmarks[indices[5]].x, landmarks[indices[5]].y]))
            p3_p5 = np.linalg.norm(np.array([landmarks[indices[2]].x, landmarks[indices[2]].y]) - np.array([landmarks[indices[4]].x, landmarks[indices[4]].y]))
            p1_p4 = np.linalg.norm(np.array([landmarks[indices[0]].x, landmarks[indices[0]].y]) - np.array([landmarks[indices[3]].x, landmarks[indices[3]].y]))
            return (p2_p6 + p3_p5) / (2.0 * p1_p4)
        except: return 0.0

    def calculate_mar(self, landmarks):
        try:
            vertical = np.linalg.norm(np.array([landmarks[13].x, landmarks[13].y]) - np.array([landmarks[14].x, landmarks[14].y]))
            horizontal = np.linalg.norm(np.array([landmarks[61].x, landmarks[61].y]) - np.array([landmarks[291].x, landmarks[291].y]))
            return vertical / horizontal
        except: return 0.0

    def get_head_pose(self, landmarks):
        try:
            nose = landmarks[1].x
            left_ear = landmarks[234].x
            right_ear = landmarks[454].x
            ratio = (nose - left_ear) / (right_ear - nose + 0.0001)
            if ratio < 0.5: return "Né tránh (Trái)"
            if ratio > 2.0: return "Né tránh (Phải)"
            if (landmarks[1].y - landmarks[10].y) < 0.03: return "Cúi đầu"
            return "Thẳng"
        except: return "KĐ"

    def detect_emotion(self, landmarks):
        try:
            mouth_corner_y = (landmarks[61].y + landmarks[291].y) / 2
            mouth_center_y = (landmarks[13].y + landmarks[14].y) / 2
            upper_lip_y = landmarks[0].y
            face_width = np.linalg.norm(np.array([landmarks[234].x, landmarks[234].y]) - np.array([landmarks[454].x, landmarks[454].y]))
            brow_dist = np.linalg.norm(np.array([landmarks[107].x, landmarks[107].y]) - np.array([landmarks[336].x, landmarks[336].y]))
            
            if mouth_corner_y < upper_lip_y: return "Tích cực/Vui vẻ 😊"
            elif mouth_corner_y > mouth_center_y + 0.0025: return "Buồn / Chán nản 😞"
            elif (brow_dist / face_width) < 0.16: return "Căng thẳng/Stress 😖"
            return "Bình thường"
        except: return "Bình thường"

    def process_zone(self, frame_crop):
        if frame_crop is None or frame_crop.size == 0: return "NO DATA", (100, 100, 100)
        rgb_crop = cv2.cvtColor(frame_crop, cv2.COLOR_BGR2RGB)
        results = self.face_mesh.process(rgb_crop)

        if results.multi_face_landmarks:
            landmarks = results.multi_face_landmarks[0].landmark
            LEFT_EYE = [33, 160, 158, 133, 153, 144]
            RIGHT_EYE = [362, 385, 387, 263, 373, 380]
            ear = (self.calculate_ear(landmarks, LEFT_EYE) + self.calculate_ear(landmarks, RIGHT_EYE)) / 2.0
            mar = self.calculate_mar(landmarks)
            pose = self.get_head_pose(landmarks)
            emotion = self.detect_emotion(landmarks)

            status = "Ổn định"
            color = (0, 255, 0)
            if ear < 0.20: status, color = "Kiệt sức / Buồn ngủ 😴", (0, 0, 255)
            elif mar > 0.5: status, color = "Mệt mỏi / Thiếu oxy 🥱", (0, 165, 255)
            elif pose == "Cúi đầu": status, color = "Thu mình / Trầm tư 🙇", (255, 0, 255)
            elif "Né tránh" in pose: status, color = f"Mất tập trung ({pose})", (0, 255, 255)
            elif "Căng thẳng" in emotion: status, color = "Căng thẳng / Lo âu 😖", (128, 0, 128)
            elif "Buồn" in emotion: status, color = "Buồn / Chán nản 😞", (0, 100, 255)
            elif "Tích cực" in emotion: status, color = "Tích cực / Hứng thú 😄", (0, 255, 127)
            return status, color
        else: return "Vắng / K.Thấy Mặt", (128, 128, 128)

# --- VIDEO THREAD ---
class VideoThread(QThread):
    change_pixmap_signal = pyqtSignal(QImage)
    update_board_signal = pyqtSignal(list)
    error_signal = pyqtSignal(str)
    
    def __init__(self, camera_index=0):
        super().__init__()
        self.camera_index = camera_index
        self._is_running = True
        self.is_monitoring = False
        self.ai = AIProcessor()
        self.tracker = BehaviorTracker(threshold_seconds=5) 
        self.zones = [] 
        self.current_drawing_zone = None
        self.frame_count = 0 
        self.last_statuses = {} 
        self.current_class = "Unknown Class"
        self.SKIP_FRAMES = 30
        
        # [MỚI] Dữ liệu điểm phạt (Score)
        self.session_risk_score = {} # {'ID': score}

    def reset_session_data(self):
        """Reset dữ liệu khi bắt đầu phiên quét mới"""
        self.session_risk_score = {}
        self.tracker.tracking_data.clear()

    # [MỚI] Hàm tính trọng số hành vi
    def get_behavior_weight(self, behavior):
        if "Ngủ" in behavior or "Kiệt sức" in behavior: return 2.0  # Phạt nặng
        if "Căng thẳng" in behavior or "Lo âu" in behavior: return 1.5 # Phạt vừa
        if "Mất tập trung" in behavior or "Né tránh" in behavior: return 1.0 # Phạt nhẹ
        return 0.0 # Bình thường / Tích cực

    def run(self):
        if sys.platform == 'win32': cap = cv2.VideoCapture(self.camera_index, cv2.CAP_DSHOW)
        else: cap = cv2.VideoCapture(self.camera_index)
        
        cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
        cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)

        if not cap.isOpened():
            self.error_signal.emit("Không thể mở Camera!")
            return

        while self._is_running:
            ret, cv_img = cap.read()
            if not ret: continue
            
            self.video_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
            self.video_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
            self.frame_count += 1
            
            if self.is_monitoring:
                final_results = []
                should_process_ai = (self.frame_count % self.SKIP_FRAMES == 0)

                for zone in self.zones:
                    z_id = str(zone["id"])
                    
                    # 1. AI Process
                    if should_process_ai:
                        x, y, w, h = zone["rect"]
                        x, y = max(0, x), max(0, y)
                        w = min(w, self.video_width - x)
                        h = min(h, self.video_height - y)
                        
                        if w > 0 and h > 0:
                            roi = cv_img[y:y+h, x:x+w]
                            status, color = self.ai.process_zone(roi)
                            self.last_statuses[z_id] = {"status": status, "color": color}
                            
                            # Cập nhật Tracker
                            report = self.tracker.update(z_id, status)
                            if report:
                                send_to_server(report, self.current_class, "epd_distraction")
                                
                                # [MỚI] Tính điểm nghiêm trọng = Thời gian * Trọng số
                                weight = self.get_behavior_weight(report['behavior'])
                                score = report['duration'] * weight
                                self.session_risk_score[z_id] = self.session_risk_score.get(z_id, 0) + score

                    # 2. Drawing & Status Update
                    cached = self.last_statuses.get(z_id, {"status": "Đang tải...", "color": (200,200,200)})
                    final_results.append({"id": z_id, "status": cached["status"], "color": cached["color"]})
                    
                    cv2.rectangle(cv_img, (zone["rect"][0], zone["rect"][1]), 
                                  (zone["rect"][0]+zone["rect"][2], zone["rect"][1]+zone["rect"][3]), 
                                  cached["color"], 2)
                    cv2.putText(cv_img, f"HS-{z_id}", (zone["rect"][0], zone["rect"][1]-5),
                                cv2.FONT_HERSHEY_SIMPLEX, 0.5, cached["color"], 1)
                
                self.update_board_signal.emit(final_results)

            else:
                # Chế độ chờ (Vẽ khung vàng)
                # Vẫn emit signal để giữ các thẻ HS hiển thị trên UI (trạng thái chờ)
                waiting_results = []
                for zone in self.zones:
                     x, y, w, h = zone["rect"]
                     z_id = str(zone['id'])
                     cv2.rectangle(cv_img, (x, y), (x+w, y+h), (0, 255, 255), 2)
                     cv2.putText(cv_img, f"HS-{z_id}", (x + 5, y - 5), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 255, 255), 2)
                     
                     # Thêm vào danh sách hiển thị UI
                     waiting_results.append({"id": z_id, "status": "Sẵn sàng", "color": (100, 100, 100)})

                if self.current_drawing_zone:
                    dx, dy, dw, dh = self.current_drawing_zone
                    cv2.rectangle(cv_img, (dx, dy), (dx+dw, dy+dh), (0, 165, 255), 2)
                
                # [MỚI] Gửi tín hiệu cập nhật UI ngay cả khi chưa Start
                if waiting_results:
                    self.update_board_signal.emit(waiting_results)

            # Convert to Qt Image
            rgb_image = cv2.cvtColor(cv_img, cv2.COLOR_BGR2RGB)
            h, w, ch = rgb_image.shape
            qt_image = QImage(rgb_image.data, w, h, ch * w, QImage.Format.Format_RGB888)
            self.change_pixmap_signal.emit(qt_image.copy())

        cap.release()

    def stop(self):
        self._is_running = False
        if self.is_monitoring:
            print("Đang tổng kết dữ liệu...")
            final_reports = self.tracker.finalize_all()
            for rep in final_reports:
                send_to_server(rep, self.current_class, "epd_distraction")
                # [MỚI] Cộng nốt điểm cuối cùng
                w = self.get_behavior_weight(rep['behavior'])
                self.session_risk_score[rep['student_id']] = self.session_risk_score.get(rep['student_id'], 0) + (rep['duration'] * w)
        self.wait()
# --- MAIN UI ---
class CameraMonitorUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Hệ Thống Giám Sát Lớp Học - AI Smart Monitor")
        self.resize(1400, 850)
        self.thread = None
        self.is_drawing_mode = False
        self.start_point = None
        self.current_temp_zone = None 
        self.student_widgets = {} 
        self.is_monitoring_active = False 
        self.current_pixmap = None
        
        # [MỚI] Biến theo dõi thời gian phiên
        self.session_start_time = None
        self.session_end_time = None

        self.setup_ui()
        self.apply_styles()
        self.load_available_cameras()
        self.session_start_time = None
        self.session_end_time = None
        self.session_snapshot_path = "session_snapshot_temp.png" # [MỚI]

        self.setup_ui()
        self.apply_styles()
        self.load_available_cameras()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.refresh_video_view()

    def refresh_video_view(self):
        if self.current_pixmap is None: return
        scaled = self.current_pixmap.scaled(self.video_label.size(), Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        self.video_label.setPixmap(scaled)

    def update_video_frame(self, image: QImage):
        self.current_pixmap = QPixmap.fromImage(image)
        self.refresh_video_view()

    def setup_ui(self):
        self.setFont(QFont("Segoe UI", 10))
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        splitter = QSplitter(Qt.Orientation.Horizontal)

        self.video_container = QWidget()
        v_layout = QVBoxLayout(self.video_container)
        v_layout.setContentsMargins(0,0,0,0)
        self.video_label = QLabel("Vui lòng chọn Camera...")
        self.video_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.video_label.setScaledContents(False) 
        self.video_label.setMouseTracking(True) 
        self.video_label.mousePressEvent = self.on_mouse_press
        self.video_label.mouseMoveEvent = self.on_mouse_move
        self.video_label.mouseReleaseEvent = self.on_mouse_release
        v_layout.addWidget(self.video_label)

        sidebar = QFrame()
        sidebar.setFixedWidth(450)
        sb_layout = QVBoxLayout(sidebar)
        
        self.grp_setup = QFrame()
        setup_layout = QVBoxLayout(self.grp_setup)
        self.combo_cam = QComboBox()
        self.combo_cam.currentIndexChanged.connect(self.start_camera_stream)
        setup_layout.addWidget(QLabel("Nguồn Camera:"))
        setup_layout.addWidget(self.combo_cam)
        setup_layout.addWidget(QLabel("🏫 Chọn Lớp Học:"))
        self.combo_class_select = QComboBox()
        self.combo_class_select.addItems(["Lớp 12A1", "Lớp 12A2", "Lớp 12A3", "Lớp 12A4", "Lớp 12A5"])
        self.combo_class_select.currentTextChanged.connect(self.update_class_name)
        setup_layout.addWidget(self.combo_class_select)
        self.btn_draw = QPushButton("🖌 THÊM VỊ TRÍ HỌC SINH")
        self.btn_draw.setCheckable(True)
        self.btn_draw.clicked.connect(self.toggle_drawing_mode)
        setup_layout.addWidget(self.btn_draw)
        self.btn_confirm_zone = QPushButton("✅ XÁC NHẬN VÙNG")
        self.btn_confirm_zone.clicked.connect(self.confirm_current_zone)
        self.btn_confirm_zone.setEnabled(False)
        setup_layout.addWidget(self.btn_confirm_zone)
        h_file = QHBoxLayout()
        btn_save = QPushButton("💾 Lưu Cấu Hình"); btn_save.clicked.connect(self.save_layout_to_file)
        btn_load = QPushButton("📂 Mở Cấu Hình"); btn_load.clicked.connect(self.load_layout_from_file)
        h_file.addWidget(btn_save); h_file.addWidget(btn_load)
        setup_layout.addLayout(h_file)
        sb_layout.addWidget(self.grp_setup)
        
        self.btn_start_monitor = QPushButton("▶ BẮT ĐẦU PHÂN TÍCH AI")
        self.btn_start_monitor.setMinimumHeight(50)
        self.btn_start_monitor.setStyleSheet("background-color: #2e7d32; font-size: 14px; font-weight: bold;")
        self.btn_start_monitor.clicked.connect(self.toggle_monitoring)
        sb_layout.addWidget(self.btn_start_monitor)

        # --- [MỚI] NÚT XUẤT BÁO CÁO WORD ---
        self.btn_export_word = QPushButton("📄 XUẤT BÁO CÁO WORD")
        self.btn_export_word.setMinimumHeight(40)
        self.btn_export_word.setStyleSheet("background-color: #0277bd; font-size: 14px; font-weight: bold;")
        self.btn_export_word.clicked.connect(self.export_word_report)
        self.btn_export_word.setEnabled(False) # Chỉ bật khi đã dừng quét
        sb_layout.addWidget(self.btn_export_word)
        
        sb_layout.addWidget(self.create_line())
        sb_layout.addWidget(QLabel("📊 TRẠNG THÁI HỌC SINH"))
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_content = QWidget()
        self.grid_layout = QGridLayout(self.scroll_content)
        self.grid_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scroll_area.setWidget(self.scroll_content)
        sb_layout.addWidget(self.scroll_area)

        splitter.addWidget(self.video_container)
        splitter.addWidget(sidebar)
        splitter.setSizes([950, 450])
        main_layout.addWidget(splitter)

    def update_class_name(self, text):
        if self.thread:
            self.thread.current_class = text
    
    def handle_thread_error(self, err_msg):
        QMessageBox.critical(self, "Lỗi", err_msg)
        self.combo_cam.setCurrentIndex(-1)

    def start_camera_stream(self):
        idx = self.combo_cam.currentIndex()
        if idx < 0: return
        if self.thread: self.thread.stop()
        self.thread = VideoThread(camera_index=idx)
        self.thread.current_class = self.combo_class_select.currentText() 
        self.thread.change_pixmap_signal.connect(self.update_video_frame)
        self.thread.update_board_signal.connect(self.update_student_panel)
        self.thread.error_signal.connect(self.handle_thread_error)
        self.thread.start()

    def toggle_monitoring(self):
        if not self.thread: return
        self.is_monitoring_active = not self.is_monitoring_active
        if self.is_monitoring_active:
            # --- BẮT ĐẦU ---
            self.thread.reset_session_data()
            self.session_start_time = datetime.datetime.now()
            self.session_end_time = None
            
            # [MỚI] CHỤP ẢNH MÀN HÌNH NGAY LẬP TỨC
            if self.current_pixmap:
                self.current_pixmap.save(self.session_snapshot_path)
            
            self.thread.is_monitoring = True
            self.btn_start_monitor.setText("⏹ DỪNG PHÂN TÍCH")
            self.btn_start_monitor.setStyleSheet("background-color: #c62828; font-size: 14px; font-weight: bold;")
            self.grp_setup.setEnabled(False)
            self.btn_export_word.setEnabled(False)
            self.is_drawing_mode = False; self.btn_draw.setChecked(False)
        else:
            # --- DỪNG ---
            self.session_end_time = datetime.datetime.now()
            
            self.thread.is_monitoring = False
            self.btn_start_monitor.setText("▶ BẮT ĐẦU PHÂN TÍCH AI")
            self.btn_start_monitor.setStyleSheet("background-color: #2e7d32; font-size: 14px; font-weight: bold;")
            self.grp_setup.setEnabled(True)
            self.btn_export_word.setEnabled(True)
            
            # Chuyển về trạng thái chờ trên giao diện
            for w in self.student_widgets.values():
                w.lbl_stat.setText("Sẵn sàng")
                w.setStyleSheet("background-color: #424242; border-radius: 6px; border: 1px solid #666;")
            
            self.thread.tracker.finalize_all()
# --- [CẬP NHẬT] HÀM XUẤT BÁO CÁO WORD (Có ngày giờ trong tên file) ---
    def export_word_report(self):
        if not self.thread or not self.session_start_time:
            QMessageBox.warning(self, "Lỗi", "Chưa có dữ liệu phiên quét!")
            return

        # Tạo tên file
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        class_name_safe = self.combo_class_select.currentText().replace(" ", "_")
        default_filename = f"BaoCao_{class_name_safe}_{timestamp}.docx"

        filename, _ = QFileDialog.getSaveFileName(self, "Lưu Báo Cáo", default_filename, "Word Documents (*.docx)")
        if not filename: return

        try:
            doc = Document()
            
            # Tiêu đề
            heading = doc.add_heading('BÁO CÁO GIÁM SÁT LỚP HỌC', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Tính toán thời gian
            end_time = self.session_end_time if self.session_end_time else datetime.datetime.now()
            duration = end_time - self.session_start_time
            duration_str = str(duration).split('.')[0]
            
            # [MỚI] Lấy ngày hiện tại
            current_date_str = datetime.datetime.now().strftime("Ngày %d tháng %m năm %Y")

            # Ghi thông tin chung
            doc.add_paragraph(f"Lớp học: {self.combo_class_select.currentText()}")
            doc.add_paragraph(f"Ngày xuất báo cáo: {current_date_str}")  # <--- Đã thêm dòng này
            doc.add_paragraph(f"Thời gian bắt đầu: {self.session_start_time.strftime('%H:%M:%S')}")
            doc.add_paragraph(f"Thời gian kết thúc: {end_time.strftime('%H:%M:%S')}")
            doc.add_paragraph(f"Tổng thời gian chạy: {duration_str}")
            
            # 1. Hình ảnh (Lấy ảnh đã chụp lúc bấm Start)
            doc.add_heading('1. Hình ảnh đầu phiên (Snapshot)', level=1)
            if os.path.exists(self.session_snapshot_path):
                doc.add_picture(self.session_snapshot_path, width=Inches(6))
            else:
                doc.add_paragraph("[Không có hình ảnh ghi nhận]")
            
            # 2. Bảng xếp hạng nghiêm trọng
            doc.add_heading('2. Xếp hạng mức độ nghiêm trọng', level=1)
            doc.add_paragraph('Điểm số (Score) = Thời gian vi phạm x Hệ số (Ngủ: 2.0, Stress: 1.5, Mất tập trung: 1.0)')
            
            # Lấy dữ liệu điểm số
            risk_data = self.thread.session_risk_score
            sorted_risk = sorted(risk_data.items(), key=lambda x: x[1], reverse=True)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            headers = ['Hạng', 'Học Sinh', 'Điểm Nguy cơ', 'Đánh Giá']
            for i, h in enumerate(headers):
                run = table.rows[0].cells[i].paragraphs[0].add_run(h)
                run.bold = True

            if not sorted_risk:
                doc.add_paragraph("✅ Lớp học rất tốt, không có nguy cơ.", style='Intense Quote')
            else:
                for idx, (sid, score) in enumerate(sorted_risk):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx + 1)
                    row_cells[1].text = f"HS-{sid}"
                    row_cells[2].text = f"{score:.1f}"
                    
                    eval_run = row_cells[3].paragraphs[0].add_run()
                    if score > 40:
                        eval_run.text = "NGUY HIỂM (Báo GV)"
                        eval_run.font.color.rgb = RGBColor(255, 0, 0) # Đỏ
                        eval_run.font.bold = True
                    elif score > 15:
                        eval_run.text = "CẢNH BÁO (Nhắc nhở)"
                        eval_run.font.color.rgb = RGBColor(255, 140, 0) # Cam
                        eval_run.font.bold = True
                    else:
                        eval_run.text = "Cần chú ý"
                        eval_run.font.color.rgb = RGBColor(0, 100, 0) # Xanh

            doc.save(filename)
            QMessageBox.information(self, "Thành công", f"Đã xuất báo cáo:\n{filename}")

        except Exception as e:
            QMessageBox.critical(self, "Lỗi Xuất File", str(e))
            traceback.print_exc()
    def update_student_panel(self, data_list):
        COLUMNS = 3
        # Lọc danh sách ID hiện có
        active_ids = [str(item['id']) for item in data_list]
        
        # Xóa widget thừa
        current_widget_ids = list(self.student_widgets.keys())
        for sid in current_widget_ids:
            if sid not in active_ids:
                self.student_widgets[sid].setParent(None)
                del self.student_widgets[sid]

        for index, item in enumerate(data_list):
            sid = str(item['id'])
            status_text = item['status']
            
            # Tạo thẻ nếu chưa có
            if sid not in self.student_widgets:
                card = QFrame()
                card.setFixedSize(130, 80) 
                l = QVBoxLayout(card); l.setContentsMargins(2,2,2,2); l.setSpacing(0)
                
                lbl_id = QLabel(f"HS {sid}")
                lbl_id.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl_id.setStyleSheet("font-weight: bold; color: white; font-size: 14px;")
                
                lbl_stat = QLabel(status_text)
                lbl_stat.setAlignment(Qt.AlignmentFlag.AlignCenter)
                lbl_stat.setWordWrap(True) 
                lbl_stat.setStyleSheet("font-size: 11px; color: #ddd;")
                
                card.lbl_stat = lbl_stat
                l.addWidget(lbl_id); l.addWidget(lbl_stat)
                
                # Context Menu Xóa
                card.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
                card.customContextMenuRequested.connect(lambda pos, s=sid: self.show_context_menu(pos, s))
                
                self.grid_layout.addWidget(card, index // COLUMNS, index % COLUMNS)
                self.student_widgets[sid] = card
            
            # Cập nhật nội dung và màu sắc
            w = self.student_widgets[sid]
            w.lbl_stat.setText(status_text)
            
            # [LOGIC MÀU SẮC]
            bg_color = "#424242" # Mặc định xám đậm (Sẵn sàng)
            border = "1px solid #666"
            
            if "Sẵn sàng" in status_text: 
                bg_color = "#424242"
            elif "NGỦ" in status_text or "Kiệt sức" in status_text: 
                bg_color = "#d32f2f" # Đỏ
                border = "none"
            elif "Vắng" in status_text: 
                bg_color = "#212121" # Đen xám
            elif "Mệt mỏi" in status_text: 
                bg_color = "#ef6c00" # Cam đậm
            elif "Mất tập trung" in status_text: 
                bg_color = "#f9a825" # Vàng
            elif "Thu mình" in status_text or "Cúi" in status_text: 
                bg_color = "#7b1fa2" # Tím
            elif "Buồn" in status_text: 
                bg_color = "#1565c0" # Xanh dương
            elif "Căng thẳng" in status_text: 
                bg_color = "#6a1b9a" # Tím đậm
            elif "Ổn định" in status_text or "Tích cực" in status_text:
                bg_color = "#2e7d32" # Xanh lá

            w.setStyleSheet(f"background-color: {bg_color}; border-radius: 6px; border: {border};")
    def save_layout_to_file(self):
        if not self.thread or not self.thread.zones: return
        name, ok = QInputDialog.getText(self, "Lưu", "Nhập tên file (không dấu):")
        if ok and name:
            with open(f"{name}.json", 'w') as f: json.dump(self.thread.zones, f)
            QMessageBox.information(self, "Thành công", "Đã lưu bản đồ lớp học.")

    def load_layout_from_file(self):
        fname, _ = QFileDialog.getOpenFileName(self, "Mở", "", "JSON (*.json)")
        if not fname or not self.thread: return
        try:
            with open(fname, 'r') as f: 
                self.thread.zones = json.load(f)
                for i in reversed(range(self.grid_layout.count())): 
                    self.grid_layout.itemAt(i).widget().setParent(None)
                self.student_widgets.clear()
                QMessageBox.information(self, "Thành công", f"Đã nạp {len(self.thread.zones)} vị trí học sinh.")
        except Exception as e: QMessageBox.warning(self, "Lỗi", str(e))

    def toggle_drawing_mode(self):
        self.is_drawing_mode = self.btn_draw.isChecked()
        self.btn_draw.setText("❌ HỦY VẼ" if self.is_drawing_mode else "🖌 THÊM VỊ TRÍ HỌC SINH")
        self.btn_draw.setStyleSheet("background-color: #d32f2f;" if self.is_drawing_mode else "")

    def confirm_current_zone(self):
        if self.current_temp_zone and self.thread:
            existing_ids = [int(z['id']) for z in self.thread.zones]
            new_id = max(existing_ids) + 1 if existing_ids else 1
            self.thread.zones.append({"id": new_id, "rect": self.current_temp_zone})
            self.thread.current_drawing_zone = None
            self.current_temp_zone = None
            self.btn_confirm_zone.setEnabled(False)
            
            # [MỚI] TẠO THẺ HỌC SINH NGAY LẬP TỨC TRÊN UI
            dummy_list = [{"id": str(z['id']), "status": "Sẵn sàng", "color": (100,100,100)} for z in self.thread.zones]
            self.update_student_panel(dummy_list)
            
            print(f"✅ Đã thêm vùng mới: HS-{new_id}")

    def get_real_coords(self, qpoint):
        if not self.thread or self.current_pixmap is None: return 0, 0
        lbl_w = self.video_label.width(); lbl_h = self.video_label.height()
        vid_w = self.thread.video_width; vid_h = self.thread.video_height
        if lbl_w == 0 or lbl_h == 0: return 0, 0
        scale = min(lbl_w / vid_w, lbl_h / vid_h)
        disp_w = vid_w * scale; disp_h = vid_h * scale
        offset_x = (lbl_w - disp_w) / 2; offset_y = (lbl_h - disp_h) / 2
        x = qpoint.x() - offset_x; y = qpoint.y() - offset_y
        if x < 0 or y < 0 or x > disp_w or y > disp_h: return None, None
        return int(x / scale), int(y / scale)

    def on_mouse_press(self, event):
        if self.is_drawing_mode and event.button() == Qt.MouseButton.LeftButton:
            pt = self.get_real_coords(event.position())
            if pt[0] is not None: self.start_point = pt

    def on_mouse_move(self, event):
        if self.is_drawing_mode and self.start_point:
            cur = self.get_real_coords(event.position())
            if cur[0] is not None:
                x1, y1 = self.start_point; x2, y2 = cur
                self.thread.current_drawing_zone = (min(x1, x2), min(y1, y2), abs(x2 - x1), abs(y2 - y1))

    def on_mouse_release(self, event):
        if self.is_drawing_mode:
            if self.thread.current_drawing_zone and self.thread.current_drawing_zone[2] > 10:
                self.current_temp_zone = self.thread.current_drawing_zone
                self.btn_confirm_zone.setEnabled(True)
            self.start_point = None

    def load_available_cameras(self):
        self.combo_cam.clear()
        for i in range(len(QMediaDevices.videoInputs())): self.combo_cam.addItem(f"Camera {i}")

    def create_line(self):
        l = QFrame(); l.setFrameShape(QFrame.Shape.HLine); l.setStyleSheet("color: #555;"); return l

    def apply_styles(self):
        self.setStyleSheet("""
            QMainWindow { background-color: #1e1e1e; color: #fff; }
            QPushButton { padding: 8px; border-radius: 4px; background: #424242; color: white; border: 1px solid #555; }
            QPushButton:hover { background: #505050; }
            QScrollArea { border: none; background: transparent; }
            QWidget { background: transparent; }
            QComboBox { padding: 5px; background: #333; color: white; border: 1px solid #555; }
        """)

    def closeEvent(self, event):
        if self.thread: self.thread.stop()
        event.accept()

    def show_context_menu(self, pos, student_id):
        menu = QMenu(self)
        delete_action = menu.addAction(f"❌ Xóa vị trí HS-{student_id}")
        sender_widget = self.student_widgets.get(student_id)
        if sender_widget:
            action = menu.exec(sender_widget.mapToGlobal(pos))
            if action == delete_action: self.delete_zone(student_id)

    def delete_zone(self, student_id):
        if not self.thread: return
        reply = QMessageBox.question(self, 'Xác nhận', f"Bạn có chắc muốn xóa vùng theo dõi HS-{student_id}?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.thread.zones = [z for z in self.thread.zones if str(z['id']) != str(student_id)]
            if student_id in self.student_widgets:
                self.student_widgets[student_id].setParent(None)
                del self.student_widgets[student_id]

def resource_path(relative_path):
    try: base_path = sys._MEIPASS
    except Exception: base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon(resource_path("AI Smart Monitor.ico")))
    window = CameraMonitorUI()
    window.show()
    sys.exit(app.exec())
